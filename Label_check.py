import streamlit as st
import pandas as pd
import re
 
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import html
import io
from PIL import Image, ImageDraw, ImageFont

@st.cache_data
def load_ins_database():
    return pd.read_csv("ins_database.csv", encoding="utf-8-sig")

@st.cache_data
def load_warnings_database():
    return pd.read_csv("warnings_database.csv", encoding="utf-8-sig")

# Define standard colors
COLOR_SUCCESS = RGBColor.from_string("006400")  # Dark Green
COLOR_FAILURE = RGBColor.from_string("8B0000")  # Dark Red
COLOR_WARNING = RGBColor.from_string("FF8C00")  # Dark Orange
COLOR_BLACK = RGBColor.from_string("000000")
COLOR_ALERT_RED = RGBColor.from_string("B91C1C")

TARGET_FONT_NAME = 'TH Sarabun New'
TARGET_FONT_SIZE = Pt(14)

ASSET_DIR = Path(__file__).parent / "assets"
GDA_IMAGE_PATH = ASSET_DIR / "gda.png"
NUTRITION_IMAGE_PATH = ASSET_DIR / "nutrition.png"
ORYOR_IMAGE_PATH = ASSET_DIR / "oryor.png"

BOXED_LABEL_TEXTS = {
    "บริโภคแต่น้อยและออกกำลังกายเพื่อสุขภาพ",
}

PREVIEW_EXCLUDE_PREFIXES = [
    "ชื่ออาหาร:",
    "แสดง ข้อมูลสำหรับผู้แพ้อาหาร",
    "แสดง 'ข้อมูลสำหรับผู้แพ้อาหาร",
    "เลขสารบบอาหาร ในเครื่องหมายแสดงเลขสารบบอาหาร",
]

PREVIEW_EXCLUDE_PREFIXES_NORMALIZED = [prefix.lower() for prefix in PREVIEW_EXCLUDE_PREFIXES]

PREVIEW_POST_INGREDIENT_KEYWORDS = [
    "วัตถุเจือปนอาหาร",
    "แต่งกลิ่น",
    "แต่งรส",
    "การแต่ง",
    "ข้อมูลสำหรับผู้แพ้อาหาร",
]

FONT_CANDIDATES = [
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/tahoma.ttf",
    "arial.ttf",
    "tahoma.ttf",
]

def _load_overlay_font(size):
    for font_path in FONT_CANDIDATES:
        try:
            path_obj = Path(font_path)
            if path_obj.exists():
                return ImageFont.truetype(str(path_obj), size=size)
            return ImageFont.truetype(font_path, size=size)
        except (OSError, IOError):
            continue
    return ImageFont.load_default()

def generate_oryor_badge(reg_number):
    text = _clean_text(reg_number)
    if not text:
        return None
    if not ORYOR_IMAGE_PATH.exists():
        return None
    try:
        base_image = Image.open(ORYOR_IMAGE_PATH).convert("RGBA")
    except (OSError, IOError):
        return None
    badge = base_image.copy()
    draw = ImageDraw.Draw(badge)
    font_size = max(18, int(badge.width * 0.22))
    font = _load_overlay_font(font_size)
    text_bbox = draw.textbbox((0, 0), text, font=font)
    text_width = text_bbox[2] - text_bbox[0]
    text_height = text_bbox[3] - text_bbox[1]
    horizontal_margin = max(6, int(badge.width * 0.08))
    vertical_margin = max(4, int(badge.height * 0.07))
    x = max(horizontal_margin, (badge.width - text_width) / 2)
    y = badge.height - text_height - vertical_margin
    if y < horizontal_margin:
        y = horizontal_margin
    background_padding = int(text_height * 0.35)
    rect_left = x - background_padding
    rect_right = x + text_width + background_padding
    rect_top = y - background_padding / 2
    rect_bottom = y + text_height + background_padding / 2
    draw.rectangle(
        [rect_left, rect_top, rect_right, rect_bottom],
        fill=(255, 255, 255, 220),
    )
    draw.text((x, y), text, font=font, fill=(0, 0, 0, 255))
    return badge

def prepare_preview_image_entries(registration_number, include_gda=True, include_nutrition=True):
    entries = []
    # ยกเลิกการทำ overlay เลขสารบบอาหารบนรูปเครื่องหมาย อย. แสดงรูปมาตรฐานเท่านั้น
    if ORYOR_IMAGE_PATH.exists():
        entries.append(
            {
                "image": ORYOR_IMAGE_PATH,
                "caption": "เครื่องหมาย อย.",
                "width": Inches(1.2),
            }
        )
    if GDA_IMAGE_PATH.exists():
        if include_gda:
            entries.append(
                {
                    "image": GDA_IMAGE_PATH,
                    "caption": "ฉลาก GDA",
                    "width": Inches(1.3),
                }
            )
    if NUTRITION_IMAGE_PATH.exists() and include_nutrition:
        entries.append(
            {
                "image": NUTRITION_IMAGE_PATH,
                "caption": "ตารางโภชนาการ\nสามารถกรอกข้อมูลและดาวน์โหลดฉลากโภชนาการและฉลาก GDA ที่สมบูรณ์ได้ที่เว็ปไซต์กองอาหาร (https://fdaconnect.fda.moph.go.th/NF_GDA/)",
                "width": Inches(1.7),
            }
        )
    return entries

def set_cell_border(cell, color="FF0000", size=12):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        edge_element = tcBorders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tcBorders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), str(size))
        edge_element.set(qn("w:color"), color)
def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for margin_name, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        mar = tcMar.find(qn(f'w:{margin_name}'))
        if mar is None:
            mar = OxmlElement(f'w:{margin_name}')
            tcMar.append(mar)
        mar.set(qn('w:w'), str(value))
        mar.set(qn('w:type'), 'dxa')


MANUFACTURER_ADDRESS_PLACEHOLDER = "[กรุณากรอกที่ตั้งตามใบอนุญาต]"
MANUFACTURER_ROLE_LABELS = {
    "ผู้ผลิต": "ผลิตโดย",
    "นำเข้า": "นำเข้าโดย",
    "แบ่งบรรจุ": "แบ่งบรรจุโดย",
}

MANUFACTURER_ROLE_PLACEHOLDERS = {
    "ผู้ผลิต": "[กรุณากรอกชื่อผู้ผลิต]",
    "นำเข้า": "[กรุณากรอกชื่อผู้นำเข้า]",
    "แบ่งบรรจุ": "[กรุณากรอกชื่อผู้แบ่งบรรจุ]",
}

def _clean_text(value):
    if isinstance(value, str):
        cleaned = re.sub(r"\s+", " ", value.strip())
        return cleaned
    return ""

def format_foreign_manufacturer_section(foreign_name, foreign_country):
    name_display = _clean_text(foreign_name) or "[กรุณากรอกชื่อผู้ผลิตในต่างประเทศ]"
    country_display = _clean_text(foreign_country) or "[กรุณากรอกประเทศผู้ผลิต]"
    return f"ผลิตโดย {name_display} ประเทศ {country_display}"

def format_manufacturer_contact(role_key, name, address):
    prefix = MANUFACTURER_ROLE_LABELS.get(role_key, "ผู้ผลิต/ผู้นำเข้า")
    name_placeholder = MANUFACTURER_ROLE_PLACEHOLDERS.get(role_key, "[กรุณากรอกชื่อผู้ผลิต/ผู้นำเข้า]")
    address_placeholder = MANUFACTURER_ADDRESS_PLACEHOLDER

    name_display = _clean_text(name) or name_placeholder
    address_display = _clean_text(address) or address_placeholder

    return f"{prefix}: {name_display} {address_display}"

# Helper function to add a styled heading with numbering
def add_styled_heading(document, text, level=1, numbered=True, section_number=""):
    prefix = f"{section_number} " if numbered and section_number else ""
    heading = document.add_heading(f"{prefix}{text}", level=level)
    for run in heading.runs:
        run.font.bold = True 
    heading.paragraph_format.space_after = Pt(6)
    return heading

# Helper function to add a paragraph with specific styling
def add_styled_paragraph(document, text, bold=False, italic=False, color=COLOR_BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    return p

def add_page_numbers(document):
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run = p.add_run()
        run.element.append(fldChar1)
        run.element.append(instrText)
        run.element.append(fldChar2)

def get_net_content_placeholder(food_consistency):
    """Return a placeholder text for net content based on food consistency."""
    if food_consistency == "ของเหลว":
        return "ปริมาตรสุทธิ ….. มล./ลิตร"
    if food_consistency == "เม็ดหรือแคปซูล":
        return "ปริมาณสุทธิ ….. เม็ด/แคปซูล"
    return "น้ำหนักสุทธิ ….. กรัม/กิโลกรัม"

def build_label_preview_context(
    food_name,
    main_ingredients,
    food_consistency,
    food_type,
    manufacturer_line,
    foreign_manufacturer_line,
    food_registration_number,
    shelf_life_option,
    has_allergen,
    allergen_groups,
    maybe_allergen,
    maybe_allergen_groups,
    has_desiccant,
    ordered_labels,
    single_ingredient_only: bool = False,
):
    """Aggregate data points for the label preview box."""
    main_ingredients = main_ingredients or []
    allergen_groups = allergen_groups or []
    maybe_allergen_groups = maybe_allergen_groups or []
    ordered_labels = ordered_labels or []
    shelf_life_option = shelf_life_option or ""

    title_text = _clean_text(food_name)
    title_is_placeholder = not title_text
    title_display = title_text or "[กรุณากรอกชื่ออาหาร]"

    cleaned_ingredients = [item for item in (_clean_text(i) for i in main_ingredients) if item]
    ingredients_text = ", ".join(cleaned_ingredients)
    ingredients_placeholder = "[กรุณากรอกส่วนประกอบสำคัญ]"

    def normalize_entry(value):
        cleaned = _clean_text(value)
        return cleaned.lower() if cleaned else ""

    boxed_map = {normalize_entry(text): text for text in BOXED_LABEL_TEXTS}
    boxed_targets = set(boxed_map.keys())
    registered_entries = set()
    registered_prefixes = set()

    def register_text(value):
        normalized = normalize_entry(value)
        if normalized:
            registered_entries.add(normalized)

    def register_prefix(label):
        cleaned_label = _clean_text(label)
        if cleaned_label:
            registered_prefixes.add(cleaned_label.lower())
            register_text(cleaned_label)

    def register_line(line):
        if line["label"]:
            register_prefix(line["label"])
            register_text(f"{line['label']}: {line['value']}")
        register_text(line["value"])

    def line_exists(line):
        checks = [line["value"]]
        if line["label"]:
            checks.append(f"{line['label']}: {line['value']}")
        return any(normalize_entry(text) in registered_entries for text in checks)

    def normalize_line(line):
        line.setdefault("label", None)
        line.setdefault("is_placeholder", False)
        normalized_value = normalize_entry(line["value"])
        display_value = line.get("display_value", line["value"])
        for boxed_key in boxed_targets:
            if boxed_key and boxed_key in normalized_value:
                line["box"] = True
                line["is_placeholder"] = False
                line["display_value"] = boxed_map.get(boxed_key, display_value)
                break
        else:
            line["box"] = line.get("box", False)
            line["display_value"] = display_value
        if line.get("box") and "display_value" not in line:
            line["display_value"] = line["value"]
        return line

    core_lines = []
    post_ingredient_lines = []
    extra_lines = []

    def append_line(line, target="core"):
        line = normalize_line(line)
        if not line_exists(line):
            if target == "core":
                core_lines.append(line)
            elif target == "post_ingredient":
                post_ingredient_lines.append(line)
            else:
                extra_lines.append(line)
            register_line(line)

    def build_line(label, value, placeholder, detect_placeholder=False):
        cleaned_value = _clean_text(value)
        if cleaned_value:
            is_placeholder = False
            if detect_placeholder and "[กรุณากรอก" in cleaned_value:
                is_placeholder = True
            line = {
                "label": label,
                "value": cleaned_value,
                "is_placeholder": is_placeholder,
            }
            append_line(line)
            return
        append_line(
            {
                "label": label,
                "value": placeholder,
                "is_placeholder": True,
            }
        )

    is_supplement = food_type == "ผลิตภัณฑ์เสริมอาหาร"
    ingredient_suffix = " (พร้อมแสดงปริมาณ)" if is_supplement else " (พร้อมแสดงร้อยละของน้ำหนักโดยประมาณของส่วนประกอบ)"
    ingredient_has_value = bool(cleaned_ingredients)
    ingredient_value = ingredients_text if ingredient_has_value else ingredients_placeholder
    # หากเลือกกรณีมีส่วนประกอบอย่างเดียว ให้ไม่แสดงหัวข้อ 'ส่วนประกอบ' ในตัวอย่างฉลาก
    if not single_ingredient_only:
        append_line(
            {
                "label": "ส่วนประกอบ",
                "value": ingredient_value,
                "display_value": f"{ingredient_value}{ingredient_suffix}",
                "is_placeholder": not ingredient_has_value,
            }
        )
    append_line(
        {
            "label": None,
            "value": get_net_content_placeholder(food_consistency),
            "is_placeholder": False,
        }
    )
    registration_clean = _clean_text(food_registration_number)
    has_registration = bool(registration_clean)
    registration_value = registration_clean or "[กรุณากรอกเลขสารบบอาหาร]"
    append_line(
        {
            "label": None,
            "value": registration_value,
            "display_value": f"{registration_value} (แสดงในกรอบเครื่องหมาย อย.)",
            "is_placeholder": not has_registration,
        }
    )
    build_line(
        None,
        manufacturer_line,
        "[กรุณากรอกข้อมูลผู้รับอนุญาต]",
        detect_placeholder=True,
    )

    if foreign_manufacturer_line:
        build_line(
            None,
            foreign_manufacturer_line,
            "[กรุณากรอกข้อมูลผู้ผลิตต่างประเทศ]",
            detect_placeholder=True,
        )

    if has_desiccant:
        append_line(
            {
                "label": None,
                "value": "คำเตือน: ห้ามรับประทานซองกันชื้น",
                "is_placeholder": False,
            }
        )

    # แสดงบรรทัด "ข้อมูลสำหรับผู้แพ้อาหาร" ในตัวอย่างฉลาก เฉพาะเมื่ออยู่ในรายการ "ข้อมูลที่ต้องมีในฉลาก"
    requires_allergen_note = any(
        "ข้อมูลสำหรับผู้แพ้อาหาร" in str(lbl) for lbl in (ordered_labels or [])
    )
    if requires_allergen_note and has_allergen and allergen_groups:
        append_line(
            {
                "label": None,
                "value": f"ข้อมูลสำหรับผู้แพ้อาหาร: มี {', '.join(allergen_groups)}",
                "is_placeholder": False,
            },
            target="post_ingredient",
        )

    if requires_allergen_note and maybe_allergen and maybe_allergen_groups:
        append_line(
            {
                "label": None,
                "value": f"ข้อมูลสำหรับผู้แพ้อาหาร: อาจมี {', '.join(maybe_allergen_groups)}",
                "is_placeholder": False,
            },
            target="post_ingredient",
        )

    top_lines = []
    bottom_lines = []

    def extract_quoted_text(label_text):
        parts = label_text.split("'")
        if len(parts) >= 3:
            return _clean_text(parts[1])
        return _clean_text(label_text)

    for label in ordered_labels:
        normalized_label = normalize_entry(label)
        if not normalized_label:
            continue
        if "คำเตือน" in normalized_label and "มีกาเฟอีน" in normalized_label:
            caffeine_line = {
                "label": None,
                "value": "มีกาเฟอีน",
                "display_value": "มีกาเฟอีน",
                "is_placeholder": False,
                "box": False,
            }
            register_line(caffeine_line)
            top_lines.append(caffeine_line)
            continue
        if "คำเตือน" in normalized_label and "ตัวอักษรขนาดไม่เล็กกว่า 1.5 มม" in normalized_label:
            warning_text = extract_quoted_text(label)
            warning_line = {
                "label": None,
                "value": warning_text,
                "display_value": warning_text,
                "is_placeholder": False,
                "box": True,
            }
            register_line(warning_line)
            bottom_lines.append(warning_line)
            continue
        if "เด็กและสตรีมีครรภ์" in normalized_label:
            warning_text = extract_quoted_text(label)
            warning_line = {
                "label": None,
                "value": warning_text,
                "display_value": warning_text,
                "is_placeholder": False,
                "box": False,
            }
            register_line(warning_line)
            bottom_lines.append(warning_line)
            continue
        if "ควรกินอาหารหลากหลาย" in normalized_label:
            warning_text = extract_quoted_text(label)
            warning_line = {
                "label": None,
                "value": warning_text,
                "display_value": warning_text,
                "is_placeholder": False,
                "box": False,
            }
            register_line(warning_line)
            bottom_lines.append(warning_line)
            continue
        if "ไม่มีผลในการป้องกัน" in normalized_label:
            warning_text = extract_quoted_text(label)
            warning_line = {
                "label": None,
                "value": warning_text,
                "display_value": warning_text,
                "is_placeholder": False,
                "box": True,
            }
            register_line(warning_line)
            bottom_lines.append(warning_line)
            continue
        if "ห้ามดื่มเกินวันละ" in normalized_label:
            warning_text = extract_quoted_text(label)
            warning_line = {
                "label": None,
                "value": warning_text,
                "display_value": warning_text,
                "is_placeholder": False,
                "box": True,
                "badge_variant": "warning",
            }
            register_line(warning_line)
            bottom_lines.append(warning_line)
            continue
        if (
            normalized_label in registered_entries
            or any(normalized_label.startswith(prefix) for prefix in registered_prefixes)
        ):
            continue
        if any(normalized_label.startswith(prefix) for prefix in PREVIEW_EXCLUDE_PREFIXES_NORMALIZED):
            continue
        append_line(
            {
                "label": None,
                "value": label,
                "is_placeholder": False,
            },
            target="post_ingredient"
            if any(keyword in label for keyword in PREVIEW_POST_INGREDIENT_KEYWORDS)
            else "extra",
        )

    additive_lines = [
        line for line in post_ingredient_lines if "วัตถุเจือปนอาหาร" in line.get("value", "")
    ]
    # ตัด prefix "วัตถุเจือปนอาหาร:" ออกจากตัวอย่างฉลาก (เช่น บรรทัดสี)
    for _line in additive_lines:
        _val = str(_line.get("value", ""))
        if _val.startswith("วัตถุเจือปนอาหาร:"):
            _line["display_value"] = _val.split(":", 1)[1].strip()
    flavor_lines = [
        line
        for line in post_ingredient_lines
        if any(keyword in line.get("value", "") for keyword in ("แต่งกลิ่น", "แต่งรส", "การแต่ง"))
    ]
    allergen_lines = [
        line for line in post_ingredient_lines if "ข้อมูลสำหรับผู้แพ้อาหาร" in line.get("value", "")
    ]
    other_post_lines = [
        line
        for line in post_ingredient_lines
        if line not in additive_lines and line not in flavor_lines and line not in allergen_lines
    ]
    ordered_post_ingredient_lines = (
        additive_lines + flavor_lines + allergen_lines + other_post_lines
    )

    preview_lines = list(top_lines)
    for line in core_lines:
        preview_lines.append(line)
        if line.get("label") == "ส่วนประกอบ":
            preview_lines.extend(ordered_post_ingredient_lines)

    added_post = len(ordered_post_ingredient_lines) > 0 and any(
        line.get("label") == "ส่วนประกอบ" for line in core_lines
    )
    if not added_post:
        preview_lines.extend(ordered_post_ingredient_lines)
    preview_lines.extend(extra_lines)
    preview_lines.extend(bottom_lines)

    return {
        "title": title_display,
        "title_is_placeholder": title_is_placeholder,
        "core_lines": core_lines,
        "extra_lines": extra_lines,
        "preview_lines": preview_lines,
    }

def generate_label_word_report(report_data):
    """สร้างรายงาน Word สำหรับการตรวจสอบฉลากอาหาร"""
    document = Document()
    
    # Set default font for the document
    style = document.styles['Normal']
    font = style.font
    font.name = TARGET_FONT_NAME
    font.size = TARGET_FONT_SIZE
    style.paragraph_format.space_after = Pt(4)

    # Main Title
    title_p = document.add_paragraph()
    title_run = title_p.add_run("รายงานผลการตรวจสอบฉลากอาหาร")
    title_run.font.bold = True 
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)

    # Add disclaimer box
    disclaimer_p = document.add_paragraph()
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer_p.add_run("⚠️ คำเตือน: แอปพลิเคชันนี้เป็นตัวช่วยในการคำนวณและตรวจสอบฉลากอาหารเท่านั้น \nไม่สามารถใช้เป็นเงื่อนไขการขออนุญาต หรืออ้างอิงทางกฎหมายได้ \nโปรดปฏิบัติตามกฎหมายอย่างเคร่งครัด")
    disclaimer_run.font.bold = True
    disclaimer_run.font.color.rgb = COLOR_WARNING
    disclaimer_p.paragraph_format.space_after = Pt(18)

    # 1. ข้อมูลพื้นฐาน
    add_styled_heading(document, "ข้อมูลพื้นฐาน", level=2, section_number="1.")
    
    add_styled_paragraph(document, f"ชื่ออาหาร: {report_data.get('food_name', 'ไม่ได้ระบุ')}")
    add_styled_paragraph(document, f"ประเภทอาหาร: {report_data.get('food_type', 'ไม่ได้ระบุ')}")
    add_styled_paragraph(document, f"ลักษณะอาหาร: {report_data.get('food_consistency', 'ไม่ได้ระบุ')}")
    add_styled_paragraph(document, f"เลขสารบบอาหาร: {report_data.get('food_registration_number', 'ไม่ได้ระบุ')}")
    add_styled_paragraph(document, f"อายุเก็บรักษา: {report_data.get('shelf_life_option', 'ไม่ได้ระบุ')}")
    
    manufacturer_role = report_data.get('manufacturer_role')
    manufacturer_line = format_manufacturer_contact(
        manufacturer_role,
        report_data.get('manufacturer_name'),
        report_data.get('manufacturer_address')
    )
    add_styled_paragraph(document, manufacturer_line)
    foreign_manufacturer_line = ""
    if manufacturer_role == "นำเข้า":
        foreign_manufacturer_line = format_foreign_manufacturer_section(
            report_data.get('foreign_manufacturer_name'),
            report_data.get('foreign_manufacturer_country')
        )
        add_styled_paragraph(document, foreign_manufacturer_line)

    document.add_paragraph()

    # 2. ส่วนประกอบและวัตถุเจือปนอาหาร
    add_styled_heading(document, "ส่วนประกอบและวัตถุเจือปนอาหาร", level=2, section_number="2.")
    
    # ส่วนประกอบหลัก
    if report_data.get('main_ingredients'):
        add_styled_paragraph(document, "ส่วนประกอบหลัก:", bold=True)
        for i, ingredient in enumerate(report_data.get('main_ingredients', []), 1):
            add_styled_paragraph(document, f"{i}. {ingredient}")
    
    flavoring_items = report_data.get('flavoring_statements', [])
    if flavoring_items:
        add_styled_paragraph(document, "การแต่งกลิ่น/รส:", bold=True)
        for item in flavoring_items:
            add_styled_paragraph(document, f"- {item}")
    
    # คำเตือนจากส่วนประกอบหลัก
    if report_data.get('ingredient_warnings'):
        add_styled_paragraph(document, "คำเตือนจากส่วนประกอบหลัก:", bold=True)
        for warning in report_data.get('ingredient_warnings', []):
            add_styled_paragraph(document, f"⚠️ {warning}", color=COLOR_WARNING)
    
    # วัตถุเจือปนอาหาร
    if report_data.get('ins_results'):
        add_styled_paragraph(document, "วัตถุเจือปนอาหาร (INS):", bold=True)
        for ins_result in report_data.get('ins_results', []):
            if ins_result.get('has_special_label'):
                add_styled_paragraph(document, f"⚠️ {ins_result.get('message')}", color=COLOR_WARNING)
            else:
                add_styled_paragraph(document, f"✅ {ins_result.get('message')}", color=COLOR_SUCCESS)
    add_styled_paragraph(
        document,
        "อาจมีข้อมูลอื่นๆเพิ่มเติม เช่น ข้อแนะนำในการเก็บรักษา วิธีปรุงเพื่อรับประทาน คำเตือนอื่นๆ นอกเหนือจากที่กฎหมายกำหนด",
        italic=True,
    )
    document.add_paragraph()

    # 3. สารก่อภูมิแพ้
    add_styled_heading(document, "สารก่อภูมิแพ้", level=2, section_number="3.")
    
    has_allergen_flag = report_data.get('has_allergen')
    allergen_groups_report = report_data.get('allergen_groups', [])
    maybe_allergen_flag = report_data.get('maybe_allergen', False)
    maybe_allergen_groups_report = report_data.get('maybe_allergen_groups', [])
    allergen_in_name_flag = report_data.get('allergen_in_name', False)

    if has_allergen_flag and allergen_groups_report:
        allergen_text = ", ".join(allergen_groups_report)
        # แสดงว่ามีสารก่อภูมิแพ้
        add_styled_paragraph(document, f"มีสารก่อภูมิแพ้: {allergen_text}", color=COLOR_WARNING)
        # หากมีการระบุในชื่ออาหารแล้ว ให้คำแนะนำว่ากรณีนี้ไม่บังคับต้องแสดง
        if allergen_in_name_flag:
            add_styled_paragraph(document, "หมายเหตุ: ได้ระบุชื่อสารก่อภูมิแพ้ไว้ในชื่ออาหารแล้ว จึงไม่บังคับให้แสดง 'ข้อมูลสำหรับผู้แพ้อาหาร' สำหรับรายการนี้", italic=True)
        else:
            add_styled_paragraph(document, "คำแนะนำ: ควรแสดง 'ข้อมูลสำหรับผู้แพ้อาหาร: มี ...' ในกรอบสี่เหลี่ยม โดยใช้ถ้อยคำที่อ่านได้ชัดเจน", italic=True)
    
    if maybe_allergen_flag and maybe_allergen_groups_report:
        maybe_text = ", ".join(maybe_allergen_groups_report)
        add_styled_paragraph(document, f"อาจมีการปนเปื้อนสารก่อภูมิแพ้: {maybe_text}", color=COLOR_WARNING)
        add_styled_paragraph(document, "คำแนะนำ: ควรแสดง 'ข้อมูลสำหรับผู้แพ้อาหาร: อาจมี ...' ในกรอบสี่เหลี่ยม โดยใช้ถ้อยคำที่อ่านได้ชัดเจน", italic=True)

    if not ((has_allergen_flag and allergen_groups_report) or (maybe_allergen_flag and maybe_allergen_groups_report)):
        add_styled_paragraph(document, "ไม่มีสารก่อภูมิแพ้", color=COLOR_SUCCESS)
    
    document.add_paragraph()

    # 4. การกล่าวอ้างโภชนาการ
    add_styled_heading(document, "การกล่าวอ้างโภชนาการ", level=2, section_number="4.")
    
    if report_data.get('has_nutrition_claim'):
        add_styled_paragraph(document, "มีการกล่าวอ้างโภชนาการ", color=COLOR_WARNING)
        add_styled_paragraph(document, "หมายเหตุ: ฉลากต้องมีตารางโภชนาการด้วย")
    else:
        add_styled_paragraph(document, "ไม่มีการกล่าวอ้างโภชนาการ", color=COLOR_SUCCESS)
    
    document.add_paragraph()

    # 5. ข้อมูลที่ต้องมีในฉลาก
    add_styled_heading(document, "ข้อมูลที่ต้องมีในฉลาก", level=2, section_number="5.")
    
    required_labels = report_data.get('required_labels', [])
    if required_labels:
        for i, label in enumerate(required_labels, 1):
            add_styled_paragraph(document, f"{i}. {label}")
    else:
        add_styled_paragraph(document, "ไม่พบข้อมูลที่ต้องแสดงในฉลาก", italic=True)
    
    document.add_paragraph()

    # 6. ตัวอย่างฉลาก
    add_styled_heading(document, "ตัวอย่างฉลาก", level=2, section_number="6.")

    label_preview = build_label_preview_context(
        report_data.get('food_name'),
        report_data.get('main_ingredients', []),
        report_data.get('food_consistency'),
        report_data.get('food_type'),
        manufacturer_line,
        foreign_manufacturer_line,
        report_data.get('food_registration_number'),
        report_data.get('shelf_life_option'),
        report_data.get('has_allergen'),
        report_data.get('allergen_groups', []),
        report_data.get('maybe_allergen'),
        report_data.get('maybe_allergen_groups', []),
        report_data.get('has_desiccant'),
        required_labels,
        single_ingredient_only=report_data.get('single_ingredient_only', False),
    )

    label_table = document.add_table(rows=1, cols=2)
    label_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    label_table.style = 'Table Grid'
    label_table.autofit = True

    text_cell = label_table.rows[0].cells[0]
    image_cell = label_table.rows[0].cells[1]

    title_paragraph = text_cell.paragraphs[0]
    title_paragraph.text = ""
    title_run = title_paragraph.add_run(label_preview["title"])
    title_run.bold = True
    if label_preview["title_is_placeholder"]:
        title_run.font.italic = True
        title_run.font.color.rgb = COLOR_WARNING
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    all_preview_lines = label_preview.get("preview_lines") or (
        label_preview["core_lines"] + label_preview["extra_lines"]
    )
    for line in all_preview_lines:
        display_text = line.get("display_value", line["value"])
        if line.get("box"):
            text_cell.add_paragraph()
            inner_table = text_cell.add_table(rows=1, cols=1)
            inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            inner_table.style = 'Table Grid'
            inner_table.autofit = True
            inner_cell = inner_table.cell(0, 0)
            inner_paragraph = inner_cell.paragraphs[0]
            inner_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inner_run = inner_paragraph.add_run(display_text)
            inner_run.bold = True
            set_cell_margins(inner_cell, top=30, start=120, bottom=30, end=120)
            variant = line.get("badge_variant")
            if variant == "warning":
                inner_run.font.color.rgb = COLOR_ALERT_RED
                set_cell_border(inner_cell, color="B91C1C")
            inner_paragraph.paragraph_format.space_after = Pt(4)
            continue

        paragraph = text_cell.add_paragraph()
        if line["label"]:
            label_run = paragraph.add_run(f"{line['label']}: ")
            label_run.bold = True
            value_run = paragraph.add_run(display_text)
        else:
            value_run = paragraph.add_run(display_text)
        if line["is_placeholder"]:
            value_run.font.italic = True
            value_run.font.color.rgb = COLOR_WARNING
        paragraph.paragraph_format.space_after = Pt(2)

    image_cell.paragraphs[0].text = ""
    include_gda_image = any("ฉลาก GDA" in str(label) for label in required_labels)
    include_nutrition_image = any("ตารางโภชนาการ" in str(label) for label in required_labels)
    image_entries = prepare_preview_image_entries(
        report_data.get('food_registration_number'),
        include_gda=include_gda_image,
        include_nutrition=include_nutrition_image,
    )
    images_added = False
    used_first_paragraph = False
    for entry in image_entries:
        image_obj = entry["image"]
        caption = entry["caption"]
        width = entry["width"]
        images_added = True
        if not used_first_paragraph:
            paragraph = image_cell.paragraphs[0]
            used_first_paragraph = True
        else:
            paragraph = image_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        if isinstance(image_obj, Image.Image):
            buffer = io.BytesIO()
            image_obj.save(buffer, format="PNG")
            buffer.seek(0)
            run.add_picture(buffer, width=width)
        else:
            run.add_picture(str(image_obj), width=width)
        caption_para = image_cell.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(10)
        caption_para.paragraph_format.space_after = Pt(4)
    if not images_added:
        placeholder_para = image_cell.paragraphs[0]
        placeholder_run = placeholder_para.add_run("ยังไม่พบไฟล์ภาพในโฟลเดอร์ assets/")
        placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder_run.italic = True
        placeholder_run.font.color.rgb = COLOR_WARNING

    document.add_paragraph()

    # 7. สรุป
    add_styled_heading(document, "สรุป", level=2, section_number="7.")
    add_styled_paragraph(document, f"พบข้อมูลที่ต้องแสดงในฉลากทั้งหมด {len(required_labels)} รายการ", color=COLOR_SUCCESS)
    add_styled_paragraph(document, f"วันที่ตรวจสอบ: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    add_styled_paragraph(
        document,
        "กรุณาศึกษาประกาศกระทรวงสาธารณสุขที่เกี่ยวข้องกับประเภทอาหารของท่านเพิ่มเติม เนื่องจากอาจมีข้อความที่กำหนดให้แสดงนอกเหนือจากนี้ ได้ที่เว็ปไซต์กองอาหาร",
        italic=True,
        color=COLOR_WARNING
    )

    add_page_numbers(document)

    # Global Font Override
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = TARGET_FONT_NAME
            run.font.size = TARGET_FONT_SIZE
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = TARGET_FONT_NAME
                        run.font.size = TARGET_FONT_SIZE

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def normalize_ins(s):
    return re.sub(r"\s+", "", str(s)).lower()

def format_label_required(template, row_data):
    """Replace placeholder tokens in label format with the INS data."""
    if not isinstance(template, str):
        return ""
    def _sanitize(value):
        if value is None:
            return ""
        if isinstance(value, float) and pd.isna(value):
            return ""
        if pd.isna(value):
            return ""
        return str(value).strip()
    replacements = {
        "ins_number": _sanitize(row_data.get("ins_number", "")),
        "name_th": _sanitize(row_data.get("name_th", "")),
        "name_en": _sanitize(row_data.get("name_en", "")),
        "function_group": _sanitize(row_data.get("function_group", "")),
    }
    pattern = re.compile(r"\b(" + "|".join(re.escape(key) for key in replacements.keys()) + r")\b")

    def _replace(match):
        key = match.group(1)
        value = replacements.get(key, "")
        return value if value else match.group(0)

    formatted = pattern.sub(_replace, template).strip()

    # ใส่คำว่า INS หน้าหมายเลขในวงเล็บ หากยังไม่มี
    ins_value = replacements.get("ins_number")
    if ins_value and not re.search(r"\(\s*ins\s*" + re.escape(ins_value) + r"\s*\)", formatted, re.IGNORECASE):
        formatted = re.sub(
            r"\(\s*" + re.escape(ins_value) + r"\s*\)",
            f"(INS {ins_value})",
            formatted
        )

    return formatted

def show():
    st.title("🔍 ตรวจสอบฉลากอาหาร")
    st.markdown("กรุณากรอกข้อมูลเพื่อตรวจสอบข้อความและคำเตือนที่ต้องแสดงในฉลากอาหาร")
    
    # Initialize session state for dynamic fields
    if "main_ingredient_count" not in st.session_state:
        st.session_state.main_ingredient_count = 5
    
    if "ins_count" not in st.session_state:
        st.session_state.ins_count = 3
    
    # Function to add more fields
    def add_main_ingredient():
        st.session_state.main_ingredient_count += 1
    
    def add_ins():
        st.session_state.ins_count += 1

    # Helper: when a suggestion is clicked, fill the input
    def _set_main_ing(idx: int, val: str):
        st.session_state[f"main_ing_{idx}"] = val

    # (previous version) We don't use a dropdown; suggestions are clickable chips

    

    # 1. ชื่ออาหาร
    st.subheader("1. ชื่ออาหาร (ตามที่ขึ้นทะเบียน)")
    food_name = st.text_input("กรอกชื่ออาหาร", placeholder="เช่น ขนมปังโฮลวีท")
    
    # 2. ประเภทอาหาร/ชนิดอาหาร
    st.subheader("2. ประเภทอาหาร/ชนิดอาหาร")
    food_type = st.selectbox(
        "เลือกประเภทอาหาร",
        [
            "อื่นๆ (ที่ไม่ใช่อาหารควบคุมเฉพาะ)",
            "อาหารขบเคี้ยว ตัวอย่างเช่น มันฝรั่งทอดกรอบ ข้าวโพดอบกรอบ ข้าวเกรียบชนิดต่างๆ ถั่วลิสงส์อบปรุงรส สาหร่ายทอดอบกรอบ ปลาหมึกแผ่นอบกรอบ หมูแผ่นอบกรอบ",
            "ช็อกโกแลต และขนมหวานรสช็อกโกแลต", 
            "ผลิตภัณฑ์ขนมอบ ตัวอย่างเช่น ขนมปังกรอบ ขนมขาไก่ เวเฟอร์สอดไส้ คุกกี้ เค้ก ขนมไหว้พระจันทร์ เอแคลร์ ครัวซองท์ พายไส้ต่างๆ",
            "อาหารกึ่งสำเร็จรูป",
            "อาหารมื้อหลักที่เป็นอาหารจานเดียว ซึ่งต้องเก็บรักษาไว้ในตู้เย็นหรือตู้แช่แข็งตลอดระยะเวลาจำหน่าย",
            "เครื่องดื่มในภาชนะบรรจุที่ปิดสนิท",
            "ชาปรุงสำเร็จ ทั้งชนิดเหลวและชนิดแห้ง",
            "กาแฟปรุงสำเร็จ ทั้งชนิดเหลวและชนิดแห้ง",
            "นมปรุงแต่ง",
            "นมเปรี้ยว",
            "ผลิตภัณฑ์ของนม",
            "น้ำนมถั่วเหลือง",
            "ไอศกรีมที่อยู่ในลักษณะพร้อมบริโภค",
            "วุ้นสำเร็จรูป",
            "ผลิตภัณฑ์เสริมอาหาร",
            "ชาจากพืช"
        ]
    )
    
    if food_type != "อื่นๆ (ที่ไม่ใช่อาหารควบคุมเฉพาะ)" and food_type != "วุ้นสำเร็จรูป" and food_type != "ผลิตภัณฑ์เสริมอาหาร" and food_type != "ชาจากพืช":
        st.info("📋 **หมายเหตุ**: อาหารประเภทนี้ต้องมีฉลาก GDA และตารางโภชนาการ")
    
    if food_type == "วุ้นสำเร็จรูป":
        st.warning("⚠️ **หมายเหตุ**: ต้องแสดง 'เด็กควรบริโภคแต่น้อย' ด้วยตัวอักษรสีแดงขนาด 5 มิลลิเมตร ในกรอบพื้นสีขาว")
    
    if food_type == "ชาจากพืช":
        st.warning("⚠️ **หมายเหตุ**: กรุณาศึกษารายชื่อพืชที่อนุญาต และคำเตือนเพิ่มเติม ในประกาศกระทรวงสาธารณสุข ฉบับที่ 426")
    
    if food_type == "ผลิตภัณฑ์เสริมอาหาร":
        st.warning("⚠️ **หมายเหตุ**: ต้องแสดงคำเตือนดังต่อไปนี้:")
        st.warning("• 'คำเตือน' ด้วยตัวอักษรขนาดไม่เล็กกว่า 1.5 มม. ในกรอบสี่เหลี่ยมสีของตัวอักษรตัดกับสีของพื้นกรอบ และสีกรอบตัดกับสีของพื้นฉลาก")
        st.warning("• 'เด็กและสตรีมีครรภ์ ไม่ควรรับประทาน' ด้วยตัวอักษรที่มีขนาดเห็นได้ชัดเจน")
        st.warning("• 'ควรกินอาหารหลากหลาย ครบ 5 หมู่ ในสัดส่วนที่เหมาะสมเป็นประจำ' ด้วยตัวอักษรที่มีขนาดเห็นได้ชัดเจน")
        st.warning("• 'ไม่มีผลในการป้องกัน หรือรักษาโรค' ด้วยตัวอักษรหนาทึบ ในกรอบสี่เหลี่ยม สีของตัวอักษรตัดกับสีของพื้นกรอบ และสีของกรอบตัดกับสีของพื้นฉลาก")
    
    # 3. ลักษณะของอาหาร
    st.subheader("3. ลักษณะของอาหาร")
    
    # กำหนดตัวเลือกลักษณะอาหารตามประเภทอาหาร
    if food_type == "ผลิตภัณฑ์เสริมอาหาร":
        consistency_options = ["ของเหลว", "ของแข็ง", "เม็ดหรือแคปซูล"]
    else:
        consistency_options = ["ของเหลว", "ของแข็ง"]
    
    food_consistency = st.radio(
        "เลือกลักษณะของอาหาร",
        consistency_options
    )
    
    if food_consistency == "ของเหลว":
        st.info("📋 **หมายเหตุ**: อาหารของเหลวใช้ปริมาตรสุทธิ (เช่น มล., ลิตร)")
    elif food_consistency == "เม็ดหรือแคปซูล":
        st.info("📋 **หมายเหตุ**: ผลิตภัณฑ์เสริมอาหารเม็ดหรือแคปซูลใช้จำนวนบรรจุ (เช่น เม็ด, แคปซูล)")
    else:
        st.info("📋 **หมายเหตุ**: อาหารของแข็งใช้น้ำหนักสุทธิ (เช่น กรัม, กิโลกรัม)")
    
    # 4. ส่วนประกอบ และวัตถุเจือปนอาหาร
    st.subheader("4. ส่วนประกอบ และวัตถุเจือปนอาหาร")
    
    # Main ingredients section
    st.markdown("**ส่วนประกอบหลัก**")
    single_ingredient_only = st.checkbox(
        "มีส่วนประกอบเพียงอย่างเดียว ไม่นับรวมวัตถุเจือปนอาหารหรือวัตถุแต่งกลิ่นรสที่เป็นส่วนผสม",
        key="single_ingredient_only"
    )
    
    # Main ingredients with inline suggestions from warnings_database
    try:
        _warnings_db = load_warnings_database()
        _warning_keywords = (
            _warnings_db["keyword"].dropna().astype(str).str.strip().tolist()
            if "keyword" in _warnings_db.columns else []
        )
    except Exception:
        _warning_keywords = []

    main_ingredients = []
    for i in range(st.session_state.main_ingredient_count):
        main_ing = st.text_input(f"ส่วนประกอบหลัก {i+1}", key=f"main_ing_{i}")

        # Show suggestions when user types >= 2 chars; keep free text otherwise
        q = (main_ing or "").strip()
        if q and len(q) >= 2 and _warning_keywords:
            suggs = [kw for kw in _warning_keywords if q.lower() in kw.lower()][:8]
            if suggs:
                st.caption("ท่านหมายถึงส่วนประกอบเหล่านี้หรือไม่ หากใช่กรุณาคลิก")
                cols = st.columns(min(len(suggs), 4))
                for j, s in enumerate(suggs):
                    with cols[j % len(cols)]:
                        st.button(s, key=f"ing_suggest_{i}_{j}", on_click=_set_main_ing, args=(i, s))

        if main_ing:
            main_ingredients.append(main_ing)
    # Add button after the last main ingredient input
    st.button("+ เพิ่มส่วนประกอบหลัก", on_click=add_main_ingredient, key="add_main")
    
    st.write("")
    
    st.markdown("**การแต่งกลิ่นและรส**")
    flavoring_options = [
        ("flavor_aroma_natural", "แต่งกลิ่นธรรมชาติ"),
        ("flavor_aroma_nature_identical", "แต่งกลิ่นเลียนธรรมชาติ"),
        ("flavor_aroma_artificial", "แต่งกลิ่นสังเคราะห์"),
        ("flavor_taste_natural", "แต่งรสธรรมชาติ"),
        ("flavor_taste_nature_identical", "แต่งรสเลียนธรรมชาติ"),
    ]
    flavoring_statements = []
    for key, label in flavoring_options:
        if st.checkbox(label, key=key):
            flavoring_statements.append(label)
    
    st.write("")
    
    # กำหนดค่าเริ่มต้นสำหรับตัวแปรกาเฟอีน
    caffeine_option = None
    container_type = None
    
    # เครื่องดื่มในภาชนะบรรจุที่ปิดสนิท - ตัวเลือกเพิ่มเติม
    if food_type == "เครื่องดื่มในภาชนะบรรจุที่ปิดสนิท":
        st.markdown("**ตัวเลือกเพิ่มเติมสำหรับเครื่องดื่ม**")
        
        caffeine_option = st.radio(
            "เลือกประเภทกาเฟอีน",
            ["ไม่มีกาเฟอีน", "ใช้วัตถุแต่งกลิ่นรสที่มีกาเฟอีนตามธรรมชาติ", "ผสมกาเฟอีนรูปแบบอื่น"],
            key="caffeine_option"
        )
        
        if caffeine_option == "ใช้วัตถุแต่งกลิ่นรสที่มีกาเฟอีนตามธรรมชาติ":
            st.warning("⚠️ **หมายเหตุ**: ต้องมีคำเตือนในฉลากว่า 'มีกาเฟอีน' ด้วยตัวอักษรขนาดความสูงไม่น้อยกว่า 2 มิลลิเมตร ที่อ่านได้ชัดเจน อยู่ในบริเวณเดียวกับชื่ออาหารหรือเครื่องหมายการค้า")
        
        elif caffeine_option == "ผสมกาเฟอีนรูปแบบอื่น":
            container_type = st.text_input("ระบุภาชนะที่ใช้บรรจุ", placeholder="เช่น กระป๋อง, ขวด", key="container_type")
            if container_type:
                st.warning(f"⚠️ **หมายเหตุ**: ต้องแสดงข้อความว่า 'ห้ามดื่มเกินวันละ 2 {container_type} เพราะอาจทำให้ใจสั่น นอนไม่หลับ เด็กและสตรีมีครรภ์ไม่ควรดื่ม ผู้มีโรคประจำตัวหรือผู้ป่วยปรึกษาแพทย์ก่อน' ด้วยตัวอักษรเส้นทึบสีแดง ขนาดความสูงไม่น้อยกว่า 2 มิลลิเมตร ในกรอบสี่เหลี่ยมพื้นขาว สีของกรอบตัดกับสีของพื้นฉลาก")
    
    # INS section
    st.markdown("**วัตถุเจือปนอาหาร**")
    
    st.caption("ตัวอย่าง: 160b(ii) กรณีมีเลขโรมันต้องมีวงเล็บครอบ")
    ins_list = []
    for i in range(st.session_state.ins_count):
        ins = st.text_input(f"เลข INS {i+1}", key=f"ins_{i}")
        if ins:
            ins_list.append(ins)
    # Add button after the last INS input
    st.button("+ เพิ่มวัตถุเจือปนอาหาร", on_click=add_ins, key="add_ins")
    
    st.markdown(
        "🔗 สามารถค้นหาเลข INS ได้ที่เว็ปไซต์ [กองอาหาร (อย.)](https://alimentum.fda.moph.go.th/FDA_FOOD_MVC/Additive/Main)"
    )
    
    # 5. สารก่อภูมิแพ้
    st.subheader("5. สารก่อภูมิแพ้")

    # แสดงกลุ่มสารก่อภูมิแพ้ตามกฎหมายเพื่อให้อ่านก่อน
    st.markdown(
        """
        ประเภทหรือชนิดของอาหารซึ่งมีสารก่อภูมิแพ้ หรือสารที่ก่อภาวะภูมิไวเกิน:
        
        - ธัญพืชที่มีกลูเตน ได้แก่ ข้าวสาลี ข้าวไรย์ ข้าวบาร์เลย์ ข้าวโอ๊ต สเปลท์ หรือสายพันธุ์ลูกผสมของธัญพืชดังกล่าว และผลิตภัณฑ์จากธัญพืชที่มีกลูเตนดังกล่าว ยกเว้น (ก) กลูโคสไซรัป หรือเดกซ์โทรสที่ได้จากข้าวสาลี (ข) มอลโทเดกซ์ตริน จากข้าวสาลี (ค) กลูโคสไซรัป จากข้าวบาร์เลย์ (ง) แอลกฮอล์ที่ได้จากการกลั่นเมล็ดธัญพืช
        - สัตว์น้ำที่มีเปลือกแข็ง เช่น ปู กุ้ง กั้ง ลอบสเตอร์ เป็นต้น และผลิตภัณฑ์จากสัตว์น้ำที่มีเปลือกแข็ง
        - ไข่ และผลิตภัณฑ์จากไข่
        - ปลา และผลิตภัณฑ์จากปลา ยกเว้น เจลาตินจากปลาที่ใช้เป็นสารช่วยพาวิตามินและแคโรทีนอยด์
        - ถั่วลิสง และผลิตภัณฑ์จากถั่วลิสง
        - ถั่วเหลือง และผลิตภัณฑ์จากถั่วเหลือง ยกเว้น (ก) น้ำมันหรือไขมันจากถั่วเหลืองที่ผ่านกระบวนการทำให้บริสุทธิ์ (ข) โทโคเฟอรอลผสม, ดี-แอลฟา-โทโคเฟอรอล, หรือ ดีแอล-แอลฟา-โทโคเฟอรอล หรือ ดี-แอลฟา-โทโคเฟอริลแอซีเทต, หรือ ดีแอล-แอลฟา-โทโคเฟอริลแอซีเทต หรือ ดี-แอลฟา-โทโคเฟอริลแอซิดซักซิเนต ที่ได้จากถั่วเหลือง (ค) ไฟโตสเตอรอล และไฟโตสเตอรอลเอสเตอร์ที่ได้จากน้ำมันถั่วเหลือง (ง) สตานอลเอสเตอร์จากพืชที่ผลิตจากสเตอรอลของน้ำมันพืชที่ได้จากถั่วเหลือง
        - นม และผลิตภัณฑ์จากนม รวมถึงแลคโตส ยกเว้น แลคติทอล
        - ถั่วที่มีเปลือกแข็ง และผลิตภัณฑ์จากถั่วที่มีเปลือกแข็ง เช่น อัลมอนต์ วอลนัท พีแคน เป็นต้น
        - ซัลไฟต์ ที่มีปริมาณมากกว่าหรือเท่ากับ 10 มิลลิกรัมต่อกิโลกรัม
        - หอย หมึก และผลิตภัณฑ์จากหอย หมึก
        """
    )

    # ฟอร์มให้ติ๊ก และกรอกเองเมื่อมี/อาจมี
    has_allergen = st.checkbox("มีสารก่อภูมิแพ้ในส่วนประกอบ")
    allergen_groups = []
    allergen_in_name = False
    if has_allergen:
        allergen_text_input = st.text_input(
            "ระบุชื่อสารก่อภูมิแพ้ที่มีในส่วนประกอบ (คั่นแต่ละรายการด้วยเครื่องหมาย ,)",
            key="allergen_text_input_direct"
        )
        if allergen_text_input:
            allergen_groups = [s.strip() for s in allergen_text_input.split(",") if s.strip()]
        allergen_in_name = st.checkbox(
            "มีการระบุชื่อสารก่อภูมิแพ้ไว้ในชื่ออาหารชัดเจน (เช่น ถั่วลิสงอบกรอบ, น้ำนมโคสด)",
            key="allergen_in_name"
        )

    maybe_allergen = st.checkbox("อาจมีการปนเปื้อนสารก่อภูมิแพ้ในกระบวนการผลิต")
    maybe_allergen_groups = []
    if maybe_allergen:
        maybe_allergen_text_input = st.text_input(
            "ระบุชื่อสารก่อภูมิแพ้ที่อาจปนเปื้อน (คั่นแต่ละรายการด้วยเครื่องหมาย ,)",
            key="allergen_text_input_maybe"
        )
        if maybe_allergen_text_input:
            maybe_allergen_groups = [s.strip() for s in maybe_allergen_text_input.split(",") if s.strip()]
    
    # 6. การกล่าวอ้างโภชนาการ
    st.subheader("6. การกล่าวอ้างโภชนาการ")
    has_nutrition_claim = st.checkbox("มีการกล่าวอ้างโภชนาการ")
    
    if has_nutrition_claim:
        st.info("📋 **หมายเหตุ**: กรุณาตรวจสอบเพิ่มเติมในเมนู 'ตรวจสอบการกล่าวอ้างโภชนาการ' และฉลากต้องมีตารางโภชนาการด้วย")
    
    # 7. เลขสารบบอาหาร
    st.subheader("7. เลขสารบบอาหาร")
    food_registration_number = st.text_input(
        "กรอกเลขสารบบอาหาร (ถ้ามี)", 
        placeholder="เช่น 12-1-12345-1-0001"
    )
    
    # 8. ชื่อและที่ตั้งผู้ผลิตหรือผู้นำเข้า
    st.subheader("8. ชื่อและที่ตั้งผู้ผลิต/ผู้นำเข้า/ผู้แบ่งบรรจุ")
    manufacturer_role = st.radio(
        "เลือกสถานะของหน่วยงาน",
        list(MANUFACTURER_ROLE_LABELS.keys()),
        index=0,
        horizontal=True
    )
    st.info(f"คำแนะนำ: แสดงคำขึ้นต้นว่า '{MANUFACTURER_ROLE_LABELS[manufacturer_role]}' บนฉลาก")

    manufacturer_name = st.text_input("ชื่อผู้รับอนุญาต")
    manufacturer_address = st.text_area("ที่ตั้งตามใบอนุญาต")
    foreign_manufacturer_name = ""
    foreign_manufacturer_country = ""
    if manufacturer_role == "นำเข้า":
        foreign_manufacturer_name = st.text_input("ชื่อผู้ผลิตในต่างประเทศ (ภาษาไทย)")
        foreign_manufacturer_country = st.text_input("ประเทศผู้ผลิต")
    
    # 9. อายุของอาหาร
    st.subheader("9. อายุของอาหาร")
    shelf_life_option = st.radio(
        "เลือกอายุเก็บรักษา",
        ["ไม่เกิน 90 วัน", "เกิน 90 วัน"]
    )
    
    shelf_life_days = 90 if shelf_life_option == "ไม่เกิน 90 วัน" else 365
    
    if shelf_life_option == "ไม่เกิน 90 วัน":
        st.info("📋 **หมายเหตุ**: อายุเก็บไม่เกิน 90 วัน ต้องระบุ วัน เดือน ปี")
    else:
        st.info("📋 **หมายเหตุ**: อายุเก็บเกิน 90 วัน สามารถระบุ เดือนและปี หรือ วัน เดือน ปี")
    
    # 10. ซองวัตถุกันชื้น
    st.subheader("10. ซองวัตถุกันชื้น")
    has_desiccant = st.checkbox("มีซองวัตถุกันชื้น")
    
    if has_desiccant:
        st.warning("⚠️ **หมายเหตุ**: ต้องระบุ 'มีซองวัตถุกันชื้น' ด้วยตัวอักษรสีแดง ขนาดตัวอักษรไม่ต่ำกว่า ๓ มิลลิเมตร บนพื้นสีขาว")
    
    # ปุ่มตรวจสอบ
    st.write("")
    st.write("")
    
    if st.button("🔍 ตรวจสอบฉลากอาหาร", type="primary"):
        # สร้างรายงานผลการตรวจสอบ
        generate_label_report(
            food_name, food_type, food_consistency, main_ingredients, ins_list,
            has_allergen, allergen_groups, has_nutrition_claim, 
            food_registration_number, manufacturer_name, manufacturer_role, manufacturer_address,
            shelf_life_option, has_desiccant, caffeine_option, container_type,
            maybe_allergen, maybe_allergen_groups, allergen_in_name, single_ingredient_only,
            flavoring_statements=flavoring_statements,
            foreign_manufacturer_name=foreign_manufacturer_name,
            foreign_manufacturer_country=foreign_manufacturer_country
        )

def generate_label_report(food_name, food_type, food_consistency, main_ingredients, ins_list,
                          has_allergen, allergen_groups, has_nutrition_claim,
                          food_registration_number, manufacturer_name, manufacturer_role, manufacturer_address,
                          shelf_life_option, has_desiccant, caffeine_option=None, container_type=None,
                          maybe_allergen=False, maybe_allergen_groups=None, allergen_in_name=False,
                          single_ingredient_only=False, flavoring_statements=None,
                          foreign_manufacturer_name="",
                          foreign_manufacturer_country=""):
    """สร้างรายงานผลการตรวจสอบฉลากอาหาร"""
    if maybe_allergen_groups is None:
        maybe_allergen_groups = []
    if flavoring_statements is None:
        flavoring_statements = []
    
    st.markdown("---")
    st.markdown("## 📋 รายงานผลการตรวจสอบฉลากอาหาร")
    
    manufacturer_line = format_manufacturer_contact(manufacturer_role, manufacturer_name, manufacturer_address)
    foreign_manufacturer_line = ""
    if manufacturer_role == "นำเข้า":
        foreign_manufacturer_line = format_foreign_manufacturer_section(
            foreign_manufacturer_name,
            foreign_manufacturer_country
        )

    # ข้อมูลพื้นฐาน
    st.markdown("### 📝 ข้อมูลพื้นฐาน")
    col1, col2 = st.columns(2)
    
    with col1:
        st.write(f"**ชื่ออาหาร**: {food_name if food_name else 'ไม่ได้ระบุ'}")
        st.write(f"**ประเภทอาหาร**: {food_type if food_type else 'ไม่ได้ระบุ'}")
        st.write(f"**ลักษณะอาหาร**: {food_consistency}")
    
    with col2:
        st.write(f"**เลขสารบบอาหาร**: {food_registration_number if food_registration_number else 'ไม่ได้ระบุ'}")
        st.write(f"**อายุเก็บรักษา**: {shelf_life_option}")
    
    st.write(f"**ข้อมูลผู้รับอนุญาต**: {manufacturer_line}")
    if foreign_manufacturer_line:
        st.write(f"**ข้อมูลผู้ผลิตต่างประเทศ**: {foreign_manufacturer_line}")

    # ส่วนประกอบและวัตถุเจือปนอาหาร
    st.markdown("### 🧪 ส่วนประกอบและวัตถุเจือปนอาหาร")
    
    required_labels = []
    ins_results = []
    ingredient_warnings = []

    # แนะนำการแสดงปริมาณคาเฟอีนสำหรับชาปรุงสำเร็จ/กาแฟปรุงสำเร็จ
    if food_type in [
        "ชาปรุงสำเร็จ ทั้งชนิดเหลวและชนิดแห้ง",
        "กาแฟปรุงสำเร็จ ทั้งชนิดเหลวและชนิดแห้ง",
    ]:
        required_labels.append(
            "แสดง 'มีกาเฟอีน ....... มก./ 100 มล.' ในกรอบสี่เหลี่ยมพื้นขาว ความสูงไม่น้อยกว่า 2 มม. ที่อ่านได้ชัดเจน บริเวณเดียวกับชื่ออาหารหรือเครื่องหมายการค้า"
        )

    # ส่วนประกอบหลัก
    if main_ingredients:
        st.markdown("#### 📋 ส่วนประกอบหลัก")
        ingredients_text = ", ".join(main_ingredients)
        st.write(f"**ส่วนประกอบ**: {ingredients_text}")
        # หากไม่ได้ยกเว้น ให้ระบุรายละเอียดการแสดงส่วนประกอบที่สำคัญตามประเภทอาหาร
        if not single_ingredient_only:
            if food_type == "ผลิตภัณฑ์เสริมอาหาร":
                required_labels.append(
                    f"ส่วนประกอบที่สำคัญ: {ingredients_text} พร้อมแสดงปริมาณ โดยให้เรียงลำดับปริมาณจากมากไปน้อย"
                )
            else:
                required_labels.append(
                    f"ส่วนประกอบที่สำคัญ: {ingredients_text} พร้อมแสดงร้อยละของน้ำหนักโดยประมาณ"
                )
        
        # คำเตือนจากส่วนประกอบหลัก
        st.markdown("#### ⚠️ คำเตือนจากส่วนประกอบหลัก")
        warnings_db = load_warnings_database()
        
        for ing in main_ingredients:
            matched = warnings_db[warnings_db["keyword"].str.strip().str.lower() == ing.lower()]
            if not matched.empty:
                row = matched.iloc[0]
                warning_message = f"คำเตือนสำหรับ '{ing}': {row['warning']}"
                st.warning(f"⚠️ {warning_message}")
                required_labels.append(f"คำเตือน: {row['warning']}")
                ingredient_warnings.append(warning_message)
            else:
                st.success(f"✅ '{ing}' ไม่พบคำเตือนเฉพาะ")

    if flavoring_statements:
        st.markdown("#### 🌸 การแต่งกลิ่น/รส")
        combined_flavoring_text = ", ".join(flavoring_statements)
        st.info(f"ต้องแสดงข้อความ: '{combined_flavoring_text}' บนฉลาก")
        if combined_flavoring_text not in required_labels:
            required_labels.append(combined_flavoring_text)
    
    # วัตถุเจือปนอาหาร
    if ins_list:
        st.markdown("#### 🔍 ผลการตรวจสอบวัตถุเจือปนอาหาร (INS)")
        ins_db = load_ins_database()
        
        ins_db["normalized"] = ins_db["ins_number"].astype(str).apply(normalize_ins)
        
        for ins in ins_list:
            ins_norm = normalize_ins(ins)
            matched = ins_db[ins_db["normalized"] == ins_norm]
            if not matched.empty:
                row = matched.iloc[0]
                row_data = row.to_dict()
                ins_number_display = str(row_data.get("ins_number", "")).strip()
                name_th_display = str(row_data.get("name_th", "")).strip()
                function_group_display = str(row_data.get("function_group", "")).strip()
                label_template = row_data.get("label_required_format", "")
                label_text = format_label_required(label_template, row_data)
                if not label_text:
                    if isinstance(label_template, str):
                        label_text = label_template.strip()
                    elif pd.isna(label_template):
                        label_text = ""
                    else:
                        label_text = str(label_template).strip()
                message = (
                    f"INS {ins_number_display} คือ {name_th_display} "
                    f"({function_group_display}) | 📋 ควรแสดงข้อความในฉลากว่า: {label_text}"
                )
                st.warning(f"⚠️ {message}")
                required_labels.append(f"วัตถุเจือปนอาหาร: {label_text}")
                ins_results.append({
                    'has_special_label': True,
                    'message': message
                })
            else:
                message = f"'{ins}' ไม่มีข้อความเฉพาะ สามารถแสดง 'วัตถุเจือปนอาหาร (INS {ins},...)' ร่วมกับวัตถุเจือปนตัวอื่นๆที่ไม่มีข้อความเฉพาะได้เลย"
                st.success(f"✅ {message}")
                required_labels.append(f"วัตถุเจือปนอาหาร (INS {ins})")
                ins_results.append({
                    'has_special_label': False,
                    'message': message
                })
    
    # สารก่อภูมิแพ้
    st.markdown("### 🚨 สารก่อภูมิแพ้")
    if has_allergen and allergen_groups:
        allergen_text = ", ".join(allergen_groups)
        if allergen_in_name:
            st.info(f"ℹ️ มีสารก่อภูมิแพ้: {allergen_text} และได้ระบุไว้ในชื่ออาหารแล้ว → ไม่บังคับให้แสดง 'ข้อมูลสำหรับผู้แพ้อาหาร' สำหรับรายการนี้")
        else:
            st.warning(f"⚠️ **มีสารก่อภูมิแพ้**: {allergen_text}")
            required_labels.append(f"แสดง ข้อมูลสำหรับผู้แพ้อาหาร: มี{allergen_text} หรือแสดง 'มี {allergen_text}' ในกรอบสี่เหลี่ยม")
    if maybe_allergen and maybe_allergen_groups:
        allergen_text2 = ", ".join(maybe_allergen_groups)
        st.warning(f"⚠️ **อาจมีการปนเปื้อนสารก่อภูมิแพ้**: {allergen_text2}")
        required_labels.append(f"แสดง ข้อมูลสำหรับผู้แพ้อาหาร: อาจมี{allergen_text2} หรือแสดง 'อาจมี {allergen_text2}' ในกรอบสี่เหลี่ยม")
    if not (has_allergen and allergen_groups) and not (maybe_allergen and maybe_allergen_groups):
        st.success("✅ **ไม่มีสารก่อภูมิแพ้**")
    
    # การกล่าวอ้างโภชนาการ
    st.markdown("### 📊 การกล่าวอ้างโภชนาการ")
    # เช็คว่าจะต้องมี GDA อยู่แล้วหรือไม่ เพื่อหลีกเลี่ยงการแจ้งซ้ำเรื่อง "ตารางโภชนาการ"
    requires_gda_ui = (
        food_type != "อื่นๆ (ที่ไม่ใช่อาหารควบคุมเฉพาะ)" and
        food_type != "วุ้นสำเร็จรูป" and
        food_type != "ผลิตภัณฑ์เสริมอาหาร" and
        food_type != "ชาจากพืช"
    )
    if has_nutrition_claim:
        st.warning("⚠️ **มีการกล่าวอ้างโภชนาการ**")
        if not requires_gda_ui:
            st.info("📋 **หมายเหตุ**: ฉลากต้องมีตารางโภชนาการด้วย")
        if not any("ตารางโภชนาการ" in str(x) for x in required_labels):
            required_labels.append("ต้องแสดงตารางโภชนาการ")
    else:
        st.success("✅ **ไม่มีการกล่าวอ้างโภชนาการ**")
    
    # ตรวจสอบประเภทอาหารที่ต้องแสดงฉลาก GDA และตารางโภชนาการ
    if food_type != "อื่นๆ (ที่ไม่ใช่อาหารควบคุมเฉพาะ)" and food_type != "วุ้นสำเร็จรูป" and food_type != "ผลิตภัณฑ์เสริมอาหาร" and food_type != "ชาจากพืช":
        st.warning("⚠️ **ประเภทอาหารที่ต้องแสดงฉลาก GDA**: ต้องแสดงฉลาก GDA และตารางโภชนาการตามประกาศฯ 394")
        required_labels.append("ต้องแสดงฉลาก GDA ตามประกาศฯ 394")
        if not any("ตารางโภชนาการ" in str(x) for x in required_labels):
            required_labels.append("ต้องแสดงตารางโภชนาการ")
    
    # คำเตือนเฉพาะตามประเภทอาหาร
    st.markdown("### ⚠️ คำเตือนเฉพาะตามประเภทอาหาร")
    
    # วุ้นสำเร็จรูป
    if food_type == "วุ้นสำเร็จรูป":
        st.warning("⚠️ **วุ้นสำเร็จรูป**: ต้องแสดง 'เด็กควรบริโภคแต่น้อย' ด้วยตัวอักษรสีแดงขนาด 5 มิลลิเมตร ในกรอบพื้นสีขาว")
        required_labels.append("แสดง 'เด็กควรบริโภคแต่น้อย' ด้วยตัวอักษรสีแดงขนาด 5 มิลลิเมตร ในกรอบพื้นสีขาว")
    
    if food_type == "ชาจากพืช":
        herbal_tea_warning = "กรุณาศึกษารายชื่อพืชที่อนุญาต และคำเตือนเพิ่มเติม ในประกาศกระทรวงสาธารณสุข ฉบับที่ 426 เนื่องจากแอปนี้ไม่สามารถตรวจสอบได้"
        st.warning(f"⚠️ **ชาจากพืช**: {herbal_tea_warning}")
        required_labels.append(f"{herbal_tea_warning}")
    
    # ผลิตภัณฑ์เสริมอาหาร
    if food_type == "ผลิตภัณฑ์เสริมอาหาร":
        st.warning("⚠️ **ผลิตภัณฑ์เสริมอาหาร**: ต้องแสดงคำเตือนดังต่อไปนี้:")
        st.warning("• 'คำเตือน' ด้วยตัวอักษรขนาดไม่เล็กกว่า 1.5 มม. ในกรอบสี่เหลี่ยม")
        st.warning("• 'เด็กและสตรีมีครรภ์ ไม่ควรรับประทาน' ด้วยตัวอักษรที่มีขนาดเห็นได้ชัดเจน")
        st.warning("• 'ควรกินอาหารหลากหลาย ครบ 5 หมู่ ในสัดส่วนที่เหมาะสมเป็นประจำ' ด้วยตัวอักษรที่มีขนาดเห็นได้ชัดเจน")
        st.warning("• 'ไม่มีผลในการป้องกัน หรือรักษาโรค' ด้วยตัวอักษรหนาทึบ ในกรอบสี่เหลี่ยม")
        required_labels.append("แสดง 'คำเตือน' ด้วยตัวอักษรขนาดไม่เล็กกว่า 1.5 มม. ในกรอบสี่เหลี่ยมสีของตัวอักษรตัดกับสีของพื้นกรอบ และสีกรอบตัดกับสีของพื้นฉลาก")
        required_labels.append("แสดง 'เด็กและสตรีมีครรภ์ ไม่ควรรับประทาน' ด้วยตัวอักษรที่มีขนาดเห็นได้ชัดเจน")
        required_labels.append("แสดง 'ควรกินอาหารหลากหลาย ครบ 5 หมู่ ในสัดส่วนที่เหมาะสมเป็นประจำ' ด้วยตัวอักษรที่มีขนาดเห็นได้ชัดเจน")
        required_labels.append("แสดง 'ไม่มีผลในการป้องกัน หรือรักษาโรค' ด้วยตัวอักษรหนาทึบ ในกรอบสี่เหลี่ยม สีของตัวอักษรตัดกับสีของพื้นกรอบ และสีของกรอบตัดกับสีของพื้นฉลาก")
    
    # อาหารขบเคี้ยว ช็อกโกแลต และผลิตภัณฑ์ขนมอบ
    if food_type in ["อาหารขบเคี้ยว ตัวอย่างเช่น มันฝรั่งทอดกรอบ ข้าวโพดอบกรอบ ข้าวเกรียบชนิดต่างๆ ถั่วลิสงส์อบปรุงรส สาหร่ายทอดอบกรอบ ปลาหมึกแผ่นอบกรอบ หมูแผ่นอบกรอบ", 
                     "ช็อกโกแลต และขนมหวานรสช็อกโกแลต", 
                     "ผลิตภัณฑ์ขนมอบ ตัวอย่างเช่น ขนมปังกรอบ ขนมขาไก่ เวเฟอร์สอดไส้ คุกกี้ เค้ก ขนมไหว้พระจันทร์ เอแคลร์ ครัวซองท์ พายไส้ต่างๆ"]:
        st.warning("⚠️ **อาหารขบเคี้ยว/ช็อกโกแลต/ขนมอบ**: ต้องแสดงข้อความในกรอบสี่เหลี่ยมว่า 'บริโภคแต่น้อยและออกกำลังกายเพื่อสุขภาพ'")
        required_labels.append("ข้อความ 'บริโภคแต่น้อยและออกกำลังกายเพื่อสุขภาพ' ด้วยตัวอักษรหนาทึบ เห็นได้ชัดเจน สีของตัวอักษรตัดกับสีพื้นของกรอบ และสีของกรอบตัดกับสีพื้นฉลาก")
    
    # เครื่องดื่มในภาชนะบรรจุที่ปิดสนิท - กาเฟอีน
    if food_type == "เครื่องดื่มในภาชนะบรรจุที่ปิดสนิท":
        if caffeine_option == "ใช้วัตถุแต่งกลิ่นรสที่มีกาเฟอีนตามธรรมชาติ":
            st.warning("⚠️ **เครื่องดื่มกาเฟอีน**: ต้องมีคำเตือน 'มีกาเฟอีน' ด้วยตัวอักษรขนาดความสูงไม่น้อยกว่า 2 มิลลิเมตร")
            required_labels.append("คำเตือน 'มีกาเฟอีน' ด้วยตัวอักษรขนาดความสูงไม่น้อยกว่า 2 มิลลิเมตร ที่อ่านได้ชัดเจน อยู่ในบริเวณเดียวกับชื่ออาหารหรือเครื่องหมายการค้า")
        elif caffeine_option == "ผสมกาเฟอีนรูปแบบอื่น" and container_type:
            st.warning(f"⚠️ **เครื่องดื่มกาเฟอีน**: ต้องแสดงข้อความ 'ห้ามดื่มเกินวันละ 2 {container_type} เพราะอาจทำให้ใจสั่น นอนไม่หลับ เด็กและสตรีมีครรภ์ไม่ควรดื่ม ผู้มีโรคประจำตัวหรือผู้ป่วยปรึกษาแพทย์ก่อน'")
            required_labels.append(f"ข้อความ 'ห้ามดื่มเกินวันละ 2 {container_type} เพราะอาจทำให้ใจสั่น นอนไม่หลับ เด็กและสตรีมีครรภ์ไม่ควรดื่ม ผู้มีโรคประจำตัวหรือผู้ป่วยปรึกษาแพทย์ก่อน' ด้วยตัวอักษรเส้นทึบสีแดง ขนาดความสูงไม่น้อยกว่า 2 มิลลิเมตร ในกรอบสี่เหลี่ยมพื้นขาว สีของกรอบตัดกับสีของพื้นฉลาก")
    
    # จัดเรียงข้อความที่ต้องมีในฉลากตามลำดับที่ต้องการ
    ordered_labels = []
    
    # 1. ชื่ออาหาร
    ordered_labels.append(f"ชื่ออาหาร: {food_name if food_name else '[กรุณากรอกชื่ออาหาร]'}")
    
    # 2. เลขสารบบอาหาร
    ordered_labels.append("เลขสารบบอาหาร ในเครื่องหมายแสดงเลขสารบบอาหาร (ดาวน์โหลดได้ที่[เว็ปไซต์กองอาหาร](https://food.fda.moph.go.th/media.php?id=629151820018753536&name=No-Color.png))")
    
    # 3. ส่วนประกอบที่สำคัญ
    if main_ingredients:
        ingredients_text = ", ".join(main_ingredients)
        if not single_ingredient_only:
            if food_type == "ผลิตภัณฑ์เสริมอาหาร":
                ordered_labels.append(
                    f"ส่วนประกอบที่สำคัญ: {ingredients_text} พร้อมแสดงปริมาณ โดยให้เรียงลำดับปริมาณจากมากไปน้อย"
                )
            else:
                ordered_labels.append(
                    f"ส่วนประกอบที่สำคัญ: {ingredients_text} พร้อมแสดงร้อยละของน้ำหนักโดยประมาณ"
                )
    
    # 4. น้ำหนัก/ปริมาณ
    if food_consistency == "ของเหลว":
        ordered_labels.append("ปริมาตรสุทธิ ….. มล./ลิตร")
    elif food_consistency == "เม็ดหรือแคปซูล":
        ordered_labels.append("จำนวนบรรจุ ….. เม็ด/แคปซูล")
    else:
        ordered_labels.append("น้ำหนักสุทธิ ….. กรัม/กิโลกรัม")
    # หากไม่กรอกส่วนประกอบเลย ให้เพิ่มข้อ 4 เป็น placeholder
    if not main_ingredients:
        ordered_labels.append("ส่วนประกอบที่สำคัญ:[กรุณากรอกข้อมูล]")

    # 5. ชื่อผู้ผลิต/ผู้นำเข้า 
    ordered_labels.append(manufacturer_line)
    if foreign_manufacturer_line:
        ordered_labels.append(f"{foreign_manufacturer_line}")

    # 6. ฉลาก GDA และตารางโภชนาการ (แสดงครั้งเดียวถ้าเข้าได้หลายเงื่อนไข)
    requires_gda = (
        food_type != "อื่นๆ (ที่ไม่ใช่อาหารควบคุมเฉพาะ)"
        and food_type != "วุ้นสำเร็จรูป"
        and food_type != "ผลิตภัณฑ์เสริมอาหาร"
        and food_type != "ชาจากพืช"
    )
    if requires_gda:
        ordered_labels.append("ต้องแสดงฉลาก GDA ตามประกาศฯ 394")
    if has_nutrition_claim or requires_gda:
        if not any("ตารางโภชนาการ" in str(x) for x in ordered_labels):
            ordered_labels.append("ต้องแสดงตารางโภชนาการ")

    # 7. อื่นๆที่เหลือ (วัตถุเจือปนอาหาร, สารก่อภูมิแพ้, การกล่าวอ้างโภชนาการ, คำเตือน, ข้อมูลเพิ่มเติม)
    for label in required_labels:
        if label not in ordered_labels:
            ordered_labels.append(label)

    # 8. อายุของอาหาร
    if shelf_life_option == "ไม่เกิน 90 วัน":
        ordered_labels.append("ควรบริโภคก่อน (ระบุ วัน เดือน ปี)")
    else:
        ordered_labels.append("ควรบริโภคก่อน (ระบุ เดือน ปี หรือ วัน เดือน ปี)")

    # 9. ซองวัตถุกันชื้น
    if has_desiccant:
        ordered_labels.append("ระบุ 'มีซองวัตถุกันชื้น' ด้วยตัวอักษรสีแดง ขนาดตัวอักษรไม่ต่ำกว่า ๓ มิลลิเมตร บนพื้นสีขาว")

    # เพิ่มข้อมูลอื่นๆที่เหลือ (แต่ไม่รวมข้อความที่ซ้ำกับ GDA และตารางโภชนาการ)
    gda_labels = ["ต้องแสดงฉลาก GDA ตามประกาศฯ 394", "ต้องแสดงตารางโภชนาการ"]
    for label in required_labels:
        if label not in ordered_labels and label not in gda_labels:
            ordered_labels.append(label)

    label_preview = build_label_preview_context(
        food_name,
        main_ingredients,
        food_consistency,
        food_type,
        manufacturer_line,
        foreign_manufacturer_line,
        food_registration_number,
        shelf_life_option,
        has_allergen,
        allergen_groups,
        maybe_allergen,
        maybe_allergen_groups,
        has_desiccant,
        ordered_labels,
        single_ingredient_only=single_ingredient_only,
    )

    # ข้อมูลที่ต้องมีในฉลาก
    st.markdown("### ✅ ข้อมูลที่ต้องมีในฉลาก")
    for i, label in enumerate(ordered_labels, 1):
        st.write(f"{i}. {label}")
    st.markdown("_อาจมีข้อมูลอื่นๆเพิ่มเติม เช่น ข้อแนะนำในการเก็บรักษา วิธีปรุงเพื่อรับประทาน คำเตือนอื่นๆ นอกเหนือจากที่กฎหมายกำหนด_")

    st.markdown("### 🏷️ ตัวอย่างฉลาก")
    st.markdown(
        """
        <style>
        .label-preview-box {
            border: 2px solid #374151;
            border-radius: 12px;
            padding: 18px 22px;
            background-color: #ffffff;
            box-shadow: 0 2px 6px rgba(0, 0, 0, 0.08);
        }
        .label-preview-title {
            font-size: 1.25rem;
            font-weight: 700;
            text-align: center;
            margin-bottom: 12px;
        }
        .label-preview-label {
            font-weight: 600;
        }
        .label-preview-line {
            margin-bottom: 6px;
        }
        .label-preview-line.placeholder,
        .label-preview-title.placeholder {
            color: #d97706;
            font-style: italic;
        }
        .label-preview-badge {
            border: 2px solid #111827;
            border-radius: 6px;
            padding: 6px 12px;
            text-align: center;
            margin: 8px 0;
            font-weight: 600;
            display: inline-block;
        }
        .label-preview-badge.placeholder {
            color: #d97706;
            font-style: italic;
        }
        .label-preview-badge.warning {
            border-color: #dc2626;
            color: #dc2626;
            background-color: #ffffff;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    preview_lines_html = []
    all_preview_lines = label_preview.get("preview_lines") or (
        label_preview["core_lines"] + label_preview["extra_lines"]
    )
    for line in all_preview_lines:
        display_text = line.get("display_value", line["value"])
        if line.get("box"):
            badge_class = "label-preview-badge"
            variant = line.get("badge_variant")
            if variant:
                badge_class += f" {variant}"
            if line.get("is_placeholder"):
                badge_class += " placeholder"
            preview_lines_html.append(
                f"<div class='{badge_class}'>{html.escape(display_text)}</div>"
            )
            continue
        label_part = ""
        if line["label"]:
            label_part = f"<span class='label-preview-label'>{html.escape(line['label'])}:</span> "
        value_part = html.escape(display_text)
        line_class = "label-preview-line"
        if line["is_placeholder"]:
            line_class += " placeholder"
        preview_lines_html.append(f"<div class='{line_class}'>{label_part}{value_part}</div>")

    title_class = "label-preview-title"
    if label_preview["title_is_placeholder"]:
        title_class += " placeholder"

    preview_html = (
        f"<div class='label-preview-box'>"
        f"<div class='{title_class}'>{html.escape(label_preview['title'])}</div>"
        f"{''.join(preview_lines_html)}"
        "</div>"
    )

    col_preview, col_images = st.columns([3, 2])
    with col_preview:
        st.markdown(preview_html, unsafe_allow_html=True)
    with col_images:
        include_gda_image = any("ฉลาก GDA" in str(label) for label in ordered_labels)
        include_nutrition_image = any("ตารางโภชนาการ" in str(label) for label in ordered_labels)
        image_entries = prepare_preview_image_entries(
            food_registration_number,
            include_gda=include_gda_image,
            include_nutrition=include_nutrition_image,
        )
        if image_entries:
            for entry in image_entries:
                image_obj = entry["image"]
                caption = entry["caption"]
                if isinstance(image_obj, Image.Image):
                    st.image(image_obj, caption=caption, use_container_width=True)
                else:
                    st.image(str(image_obj), caption=caption, use_container_width=True)
        else:
            st.info("ยังไม่พบไฟล์ภาพในโฟลเดอร์ assets/")
    
    # สรุป
    st.markdown("### 📊 สรุป")
    st.success(f"✅ พบข้อมูลที่ต้องแสดงในฉลากทั้งหมด {len(ordered_labels)} รายการ")
    st.warning("กรุณาศึกษาประกาศกระทรวงสาธารณสุขที่เกี่ยวข้องกับประเภทอาหารของท่านเพิ่มเติม เนื่องจากอาจมีข้อความที่กำหนดให้แสดงนอกเหนือจากนี้ ได้ที่[เว็ปไซต์กองอาหาร](https://food.fda.moph.go.th/food-law/category/food-product/)")

    # สร้างข้อมูลสำหรับรายงาน Word
    report_data = {
        'food_name': food_name,
        'food_type': food_type,
        'food_consistency': food_consistency,
        'food_registration_number': food_registration_number,
        'manufacturer_name': manufacturer_name,
        'manufacturer_role': manufacturer_role,
        'manufacturer_address': manufacturer_address,
        'foreign_manufacturer_name': foreign_manufacturer_name,
        'foreign_manufacturer_country': foreign_manufacturer_country,
        'shelf_life_option': shelf_life_option,
        'has_allergen': has_allergen,
        'allergen_groups': allergen_groups,
        'maybe_allergen': maybe_allergen,
        'maybe_allergen_groups': maybe_allergen_groups,
        'allergen_in_name': allergen_in_name,
        'has_nutrition_claim': has_nutrition_claim,
        'main_ingredients': main_ingredients,
        'flavoring_statements': flavoring_statements,
        'ins_results': ins_results,
        'ingredient_warnings': ingredient_warnings,
        'required_labels': ordered_labels,
        'has_desiccant': has_desiccant,
        'caffeine_option': caffeine_option,
        'container_type': container_type,
        'single_ingredient_only': single_ingredient_only,
    }
    
    # ปุ่มดาวน์โหลดรายงาน
    st.markdown("### 📥 ดาวน์โหลดรายงาน")
    
    try:
        word_stream = generate_label_word_report(report_data)
        st.download_button(
            label="📥 ดาวน์โหลดรายงาน Word (.docx)",
            data=word_stream.getvalue(),
            file_name=f"label_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    except Exception as e:
        st.error(f"เกิดข้อผิดพลาดในการสร้างรายงาน Word: {str(e)}")
