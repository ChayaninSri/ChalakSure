from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import io
import re

# Define standard colors
COLOR_SUCCESS = RGBColor.from_string("006400")  # Dark Green
COLOR_FAILURE = RGBColor.from_string("8B0000")  # Dark Red
COLOR_WARNING = RGBColor.from_string("FF8C00")  # Dark Orange
COLOR_BLACK = RGBColor.from_string("000000")

TARGET_FONT_NAME = 'TH Sarabun New'
TARGET_FONT_SIZE = Pt(14)

# Helper function to add a styled heading with numbering
def add_styled_heading(document, text, level=1, numbered=True, section_number=""):
    prefix = f"{section_number} " if numbered and section_number else ""
    heading = document.add_heading(f"{prefix}{text}", level=level)
    # Font and size will be set globally later, but keep bold for headings
    for run in heading.runs:
        run.font.bold = True 
    heading.paragraph_format.space_after = Pt(6)
    return heading

# Helper function to add a paragraph with specific styling
def add_styled_paragraph(document, text, bold=False, italic=False, color=COLOR_BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = document.add_paragraph()
    run = p.add_run(text)
    # Font and size will be set globally later
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    return p

# Helper function to add a table from a DataFrame
def add_df_to_table(document, df, title=None, include_index=False):
    if title:
        p_title = document.add_paragraph()
        run_title = p_title.add_run(title)
        # Font and size will be set globally later
        run_title.font.bold = True 
        p_title.paragraph_format.space_after = Pt(4)
    
    if df.empty:
        add_styled_paragraph(document, "ไม่มีข้อมูล", italic=True)
        document.add_paragraph() # Add some space
        return

    table = document.add_table(rows=1, cols=len(df.columns) + (1 if include_index else 0))
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Header row styling
    hdr_cells = table.rows[0].cells
    if include_index:
        cell = hdr_cells[0]
        cell.text = df.index.name if df.index.name else ''
        cell.paragraphs[0].runs[0].font.bold = True
        for i, col_name in enumerate(df.columns):
            cell = hdr_cells[i+1]
            cell.text = str(col_name)
            cell.paragraphs[0].runs[0].font.bold = True
    else:
        for i, col_name in enumerate(df.columns):
            cell = hdr_cells[i]
            cell.text = str(col_name)
            cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    if include_index:
        for index_val, row_series in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(index_val)
            for i, cell_value in enumerate(row_series):
                process_cell_value(row_cells[i+1], cell_value)
    else:
        for row_array in df.values:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_array):
                process_cell_value(row_cells[i], cell_value)
    
    document.add_paragraph() # Add some space after the table

def process_cell_value(cell, value):
    """Process a cell value with special handling for text containing bullet points and newlines"""
    value_str = str(value)
    
    # Clear the cell's default paragraph
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    
    if '\n' in value_str:
        # For text with newlines, create separate paragraphs for better spacing
        paragraphs = value_str.split('\n')
        for i, para_text in enumerate(paragraphs):
            if para_text.strip():  # Skip empty paragraphs
                p = cell.add_paragraph()
                # Add some space before bullet points for better readability
                if para_text.strip().startswith('•'):
                    p.paragraph_format.left_indent = Inches(0.1)
                    p.paragraph_format.space_before = Pt(3)
                p.add_run(para_text.strip())
    else:
        # Simple case: just set the text
        cell.text = value_str

def add_page_numbers(document):
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run = p.add_run()
        run.element.append(fldChar1)
        run.element.append(instrText)
        run.element.append(fldChar2)
        # Page number font will also be set by global override


def generate_nutrition_report(report_data: dict):
    document = Document()
    
    # Set default font for the document - this is a base, will be overridden
    style = document.styles['Normal']
    font = style.font
    font.name = TARGET_FONT_NAME
    font.size = TARGET_FONT_SIZE
    style.paragraph_format.space_after = Pt(4)

    # Main Title
    title_p = document.add_paragraph()
    title_run = title_p.add_run("รายงานผลการตรวจสอบคำกล่าวอ้างทางโภชนาการ")
    # Font and size will be set globally later, but keep bold for title
    title_run.font.bold = True 
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)

    # Add disclaimer box
    disclaimer_p = document.add_paragraph()
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer_p.add_run("⚠️ คำเตือน: แอปพลิเคชันนี้เป็นตัวช่วยในการคำนวณและตรวจสอบฉลากอาหารเท่านั้น \nไม่สามารถใช้เป็นเงื่อนไขการขออนุญาต หรืออ้างอิงทางกฎหมายได้ \nโปรดปฏิบัติตามกฎหมายอย่างเคร่งครัด")
    disclaimer_run.font.bold = True
    disclaimer_run.font.color.rgb = COLOR_WARNING
    disclaimer_p.paragraph_format.space_after = Pt(18)

    # 1. User Inputs Section
    add_styled_heading(document, "ข้อมูลที่ผู้ใช้กรอก", level=2, section_number="1.")
    add_styled_paragraph(document, f"กลุ่มอาหาร: {report_data.get('selected_label', 'N/A')}")
    if report_data.get('selected_label') == "ไม่อยู่ในบัญชีหมายเลข 2":
        add_styled_paragraph(document, f"ลักษณะของอาหาร: {'ของแข็ง (กรัม)' if report_data.get('food_state_value') == 'solid' else 'ของเหลว (มิลลิลิตร)'}")
    add_styled_paragraph(document, f"วิธีการตรวจสอบ: {report_data.get('nutrition_check_method', 'N/A')}")
    
    # Add warning message for nutrition label check method
    if report_data.get('nutrition_check_method') == "ตรวจสอบจากฉลากโภชนาการ (ต่อ 1 หน่วยบริโภค)":
        warning_p = document.add_paragraph()
        warning_run = warning_p.add_run("⚠️ ทั้งนี้ การตรวจสอบข้อมูลจากฉลากโภชนาการ เป็นการตรวจสอบจากตัวเลขที่ผ่านการปัดมาแล้ว ดังนั้นอาจทำให้ผลการคำนวณคลาดเคลื่อนจากความเป็นจริง หากท่านมีผลวิเคราะห์ แนะนำให้ใช้การตรวจสอบจากผลวิเคราะห์จะมีความแม่นยำกว่า")
        warning_run.font.color.rgb = COLOR_WARNING
        warning_p.paragraph_format.space_after = Pt(8)
    
    add_styled_paragraph(document, f"ปริมาณหน่วยบริโภคที่ระบุในฉลาก: {report_data.get('actual_serving_size', 'N/A')} g/ml")
    # Reference serving size paragraph (always shown)
    unit_text = 'กรัม' if report_data.get('food_state_value') == 'solid' else 'มิลลิลิตร'
    add_styled_paragraph(document, f"ปริมาณหน่วยบริโภคอ้างอิง: {report_data.get('ref_serving_size', 'N/A')} {unit_text}")
    # Product consumption status (only for foods not in List 2)
    if not report_data.get('is_in_list_2', False):
        add_styled_paragraph(document, f"สถานะผลิตภัณฑ์: {report_data.get('prep_option', 'N/A')}")
    if report_data.get('table_type') == "table1" and report_data.get('has_added_sugar') is not None:
        add_styled_paragraph(document, f"การเติมน้ำตาล: {report_data.get('has_added_sugar', 'N/A')}")
    
    p_input_title = document.add_paragraph()
    run_input_title = p_input_title.add_run("ข้อมูลสารอาหารที่กรอก")
    run_input_title.font.bold = True 
    p_input_title.paragraph_format.space_after = Pt(4)

    nutrient_inputs = report_data.get("nutrient_inputs", {})
    if nutrient_inputs:
        input_data_for_table = []
        for key, value in nutrient_inputs.items():
            if value is not None and not key.endswith("_is_direct_rdi"):
                thai_name = report_data.get("RDI_MAPPING_ витамин", {}).get(key, key.replace("_", " ").title())
                unit_key = key.replace("_is_direct_rdi","")
                unit = "%RDI" if nutrient_inputs.get(key + "_is_direct_rdi") else report_data.get("VITAMIN_MINERAL_UNITS", {}).get(unit_key, "g/mg/µg")
                input_data_for_table.append({"สารอาหาร": thai_name, "ปริมาณ": value, "หน่วย": unit})
        if input_data_for_table:
            inputs_df = pd.DataFrame(input_data_for_table)
            add_df_to_table(document, inputs_df)
        else:
            add_styled_paragraph(document, "ไม่พบข้อมูลสารอาหารที่กรอก", italic=True)
    else:
        add_styled_paragraph(document, "ไม่พบข้อมูลสารอาหารที่กรอก", italic=True)
    document.add_paragraph()

    # 2. Calculation Results Section
    add_styled_heading(document, "ผลการคำนวณ", level=2, section_number="2.")
    adjusted_values = report_data.get("adjusted_nutrient_values", {})
    calc_section_has_data = False
    if report_data.get("is_in_list_2"):
        p_adj_title = document.add_paragraph()
        run_adj_title = p_adj_title.add_run("ค่าสารอาหารที่ปรับตามหน่วยบริโภคอ้างอิง:")
        run_adj_title.font.bold = True
        p_adj_title.paragraph_format.space_after = Pt(4)
        adj_nut_df_data = []
        for key, value in adjusted_values.items():
            if value is not None and not key.endswith(tuple(report_data.get("REPORT_IGNORE_SUFFIXES",["_rdi_percent", "_per_100kcal", "_energy_percent", "_is_direct_rdi"]))):
                thai_name = report_data.get("RDI_MAPPING_ витамин", {}).get(key, key.replace("_", " ").title())
                unit = report_data.get("VITAMIN_MINERAL_UNITS", {}).get(key, "g/mg/µg")
                adj_nut_df_data.append({"สารอาหาร": thai_name, "ปริมาณ": f"{value:.2f}", "หน่วย": unit})
        if adj_nut_df_data:
            adj_nut_df = pd.DataFrame(adj_nut_df_data)
            add_df_to_table(document, adj_nut_df)
            calc_section_has_data = True
        else:
            add_styled_paragraph(document, "ไม่มีข้อมูล", italic=True)
    elif report_data.get("nutrition_check_method") == "ตรวจสอบจากฉลากโภชนาการ (ต่อ 1 หน่วยบริโภค)" and not report_data.get("is_in_list_2"):
        p_calc100_title = document.add_paragraph()
        run_calc100_title = p_calc100_title.add_run("ค่าสารอาหารที่คำนวณต่อ 100g/ml (จากข้อมูลฉลาก):")
        run_calc100_title.font.bold = True
        p_calc100_title.paragraph_format.space_after = Pt(4)
        calc_per_100_df_data = []
        for key, value in adjusted_values.items():
            if value is not None and not key.endswith(tuple(report_data.get("REPORT_IGNORE_SUFFIXES",["_rdi_percent", "_per_100kcal", "_energy_percent", "_is_direct_rdi"]))):
                thai_name = report_data.get("RDI_MAPPING_ витамин", {}).get(key, key.replace("_", " ").title())
                unit = report_data.get("VITAMIN_MINERAL_UNITS", {}).get(key, "g/mg/µg")
                calc_per_100_df_data.append({"สารอาหาร": thai_name, "ปริมาณ": f"{value:.2f}", "หน่วย": unit})
        if calc_per_100_df_data:
            calc_per_100_df = pd.DataFrame(calc_per_100_df_data)
            add_df_to_table(document, calc_per_100_df)
            calc_section_has_data = True
        else:
            add_styled_paragraph(document, "ไม่มีข้อมูล", italic=True)
    
    rounded_df = report_data.get("rounded_values_display_df")
    if rounded_df is not None and not rounded_df.empty:
        add_df_to_table(document, rounded_df, title="ผลการปัดเลขสารอาหารตามหลักเกณฑ์กฎหมาย")
        calc_section_has_data = True

    rdi_df = report_data.get("rdi_display_df")
    if rdi_df is not None and not rdi_df.empty:
        add_df_to_table(document, rdi_df, title="%Thai RDI ที่คำนวณได้")
        calc_section_has_data = True
        
    sf_energy_percent = report_data.get("saturated_fat_energy_percent")
    if sf_energy_percent is not None:
        add_styled_paragraph(document, f"พลังงานจากไขมันอิ่มตัว (จากค่าที่ปรับตามหน่วยบริโภคอ้างอิง หรือ ต่อ 100g/ml): {sf_energy_percent:.1f}%")
        calc_section_has_data = True
    
    label_sf_energy_percent = report_data.get("label_saturated_fat_energy_percent")
    if report_data.get("is_in_list_2") and label_sf_energy_percent is not None:
        add_styled_paragraph(document, f"พลังงานจากไขมันอิ่มตัว (จากค่าบนฉลากต่อหน่วยบริโภค): {label_sf_energy_percent:.1f}%")
        calc_section_has_data = True
        
    if not calc_section_has_data:
        add_styled_paragraph(document, "ไม่มีข้อมูลการคำนวณเพิ่มเติม", italic=True)
    document.add_paragraph()

    # 3. Evaluation Results Section
    add_styled_heading(document, "ผลการประเมินคำกล่าวอ้าง", level=2, section_number="3.")
    evaluation_messages = report_data.get("evaluation_messages", [])

    if evaluation_messages:
        table_data = []
        column_names = ["สารอาหาร", "ผลการประเมิน", "เงื่อนไขการกล่าวอ้าง"]

        for msg_data in evaluation_messages:
            text = msg_data.get("text", "N/A")
            is_success = msg_data.get("is_success", False)
            conditions_field = msg_data.get("conditions_text")

            row_dict = {cn: "" for cn in column_names}
            row_dict[column_names[0]] = "N/A" 
            row_dict[column_names[1]] = text 

            match = re.match(r"([✅❌⚠️]?)\\s*([^:]+):\\s*(.*)", text, re.DOTALL)

            if match:
                emoji = match.group(1).strip()
                nutrient_name = match.group(2).strip()
                message_body = match.group(3).strip() 

                # Extract the nutrient name for column 1
                row_dict[column_names[0]] = nutrient_name

                # For column 2, we only want the message body without the nutrient name
                # First check if the message_body starts with the nutrient name again (e.g., "nutrient_name: actual message")
                nutrient_prefix_pattern = f"^{re.escape(nutrient_name)}\\s*:"
                message_body = re.sub(nutrient_prefix_pattern, "", message_body).strip()

                pinned_notes_list = []
                evaluation_part_of_message = message_body
                
                # Check if this is a no sugar added claim with embedded conditions
                no_sugar_added_conditions = None
                if "ไม่เติมน้ำตาล" in evaluation_part_of_message and "**เงื่อนไขการกล่าวอ้าง:**" in evaluation_part_of_message:
                    parts = evaluation_part_of_message.split("**เงื่อนไขการกล่าวอ้าง:**", 1)
                    if len(parts) > 1:
                        evaluation_part_of_message = parts[0].strip()
                        raw_conditions = parts[1].strip()
                        
                        # Format the conditions with bullet points and proper line breaks
                        # Replace numbered list with bullet points
                        formatted_conditions = []
                        for line in raw_conditions.split('\n'):
                            line = line.strip()
                            if re.match(r'^\d+\.', line):  # Starts with number and dot
                                # Replace the number with bullet point
                                line = re.sub(r'^\d+\.', '•', line)
                            formatted_conditions.append(line)
                        
                        # Join with proper line breaks between points
                        no_sugar_added_conditions = '\n'.join(formatted_conditions)
                
                if '\\n   📌' in evaluation_part_of_message:
                    parts = evaluation_part_of_message.split('\\n   📌', 1)
                    evaluation_part_of_message = parts[0].strip()
                    if len(parts) > 1:
                        remaining_pinned_text = parts[1]
                        # Preserve multiple pinned notes if they exist and are separated by \n   📌
                        pinned_notes_list.append("📌" + remaining_pinned_text.replace('\\n   📌', '\\n📌'))
                
                condition_in_text_match = re.search(r'(\([^)]+\))', evaluation_part_of_message)
                extracted_condition_from_text = ""
                evaluation_text_for_display = evaluation_part_of_message

                if condition_in_text_match:
                    extracted_condition_from_text = condition_in_text_match.group(1)
                
                full_evaluation_text = evaluation_text_for_display
                if pinned_notes_list:
                    full_evaluation_text += "\\n" + "\\n".join(pinned_notes_list)
                
                # Add emoji, but don't include nutrient name before the evaluation text
                row_dict[column_names[1]] = f"{emoji} {full_evaluation_text}".strip()

                if is_success:
                    col3_parts = []
                    if extracted_condition_from_text: 
                        col3_parts.append(extracted_condition_from_text)

                    # Add the no sugar added conditions if present
                    if no_sugar_added_conditions:
                        col3_parts.append(no_sugar_added_conditions)
                        
                    raw_conditions_text_from_data = msg_data.get("conditions_text")
                    # Assign to a temp variable that can be modified
                    temp_cond_field_for_col3 = raw_conditions_text_from_data 
                    
                    if temp_cond_field_for_col3 and temp_cond_field_for_col3 not in ["Warning", "N/A", ""]:
                        claim_condition_prefix = "เงื่อนไขการกล่าวอ้าง:"
                        vitamin_mineral_condition_prefix = "เงื่อนไขการกล่าวอ้างกลุ่มวิตามินและแร่ธาตุ:"
                        
                        # Check and strip vitamin/mineral specific prefix first
                        if temp_cond_field_for_col3.startswith(vitamin_mineral_condition_prefix + " "):
                            temp_cond_field_for_col3 = temp_cond_field_for_col3[len(vitamin_mineral_condition_prefix) + 1:].strip()
                        elif temp_cond_field_for_col3.startswith(vitamin_mineral_condition_prefix):
                            temp_cond_field_for_col3 = temp_cond_field_for_col3[len(vitamin_mineral_condition_prefix):].strip()
                        # Then check and strip general claim condition prefix
                        elif temp_cond_field_for_col3.startswith(claim_condition_prefix + " "):
                            temp_cond_field_for_col3 = temp_cond_field_for_col3[len(claim_condition_prefix) + 1:].strip()
                        elif temp_cond_field_for_col3.startswith(claim_condition_prefix):
                            temp_cond_field_for_col3 = temp_cond_field_for_col3[len(claim_condition_prefix):].strip()
                        
                        # Add to parts if not empty after stripping
                        if temp_cond_field_for_col3: # Check if it's a non-empty string
                            col3_parts.append(temp_cond_field_for_col3)
                    
                    row_dict[column_names[2]] = "\\n".join(col3_parts).strip()
            else:
                # Fallback for messages not matching the primary regex structure
                first_colon_idx = text.find(':')
                if first_colon_idx != -1:
                    nutrient_name = text[:first_colon_idx].replace("✅", "").replace("❌", "").replace("⚠️", "").strip()
                    row_dict[column_names[0]] = nutrient_name
                    
                    # Extract the message after the colon for column 2, including any emoji from the original text
                    message_after_colon = text[first_colon_idx + 1:].strip()
                    emoji = ""
                    if any(e in text for e in ["✅", "❌", "⚠️"]):
                        for e in ["✅", "❌", "⚠️"]:
                            if e in text:
                                emoji = e
                                break
                    
                    # Check if this is a no sugar added claim with embedded conditions
                    no_sugar_added_conditions = None
                    if "ไม่เติมน้ำตาล" in message_after_colon and "**เงื่อนไขการกล่าวอ้าง:**" in message_after_colon:
                        parts = message_after_colon.split("**เงื่อนไขการกล่าวอ้าง:**", 1)
                        if len(parts) > 1:
                            message_after_colon = parts[0].strip()
                            raw_conditions = parts[1].strip()
                            
                            # Format the conditions with bullet points and proper line breaks
                            # Replace numbered list with bullet points
                            formatted_conditions = []
                            for line in raw_conditions.split('\n'):
                                line = line.strip()
                                if re.match(r'^\d+\.', line):  # Starts with number and dot
                                    # Replace the number with bullet point
                                    line = re.sub(r'^\d+\.', '•', line)
                                formatted_conditions.append(line)
                            
                            # Join with proper line breaks between points
                            no_sugar_added_conditions = '\n'.join(formatted_conditions)
                    
                    # For column 2, remove any possible repetition of the nutrient name at the beginning
                    nutrient_prefix_pattern = f"^{re.escape(nutrient_name)}\\s*:"
                    message_after_colon = re.sub(nutrient_prefix_pattern, "", message_after_colon).strip()
                    
                    row_dict[column_names[1]] = f"{emoji} {message_after_colon}".strip()
                    
                    # Add no sugar added conditions to column 3 if present
                    if no_sugar_added_conditions and is_success:
                        row_dict[column_names[2]] = no_sugar_added_conditions
                # else: row_dict[column_names[0]] remains "N/A", row_dict[column_names[1]] is `text`
                
                # Apply prefix stripping for column 3 in this fallback case as well
                if is_success and conditions_field and conditions_field not in ["Warning", "N/A", ""] and not no_sugar_added_conditions:
                    processed_fallback_conditions = conditions_field
                    claim_condition_prefix = "เงื่อนไขการกล่าวอ้าง:"
                    vitamin_mineral_condition_prefix = "เงื่อนไขการกล่าวอ้างกลุ่มวิตามินและแร่ธาตุ:"
                    
                    # Check and strip vitamin/mineral specific prefix first
                    if processed_fallback_conditions.startswith(vitamin_mineral_condition_prefix + " "):
                        processed_fallback_conditions = processed_fallback_conditions[len(vitamin_mineral_condition_prefix) + 1:].strip()
                    elif processed_fallback_conditions.startswith(vitamin_mineral_condition_prefix):
                        processed_fallback_conditions = processed_fallback_conditions[len(vitamin_mineral_condition_prefix):].strip()
                    # Then check and strip general claim condition prefix
                    elif processed_fallback_conditions.startswith(claim_condition_prefix + " "):
                        processed_fallback_conditions = processed_fallback_conditions[len(claim_condition_prefix) + 1:].strip()
                    elif processed_fallback_conditions.startswith(claim_condition_prefix):
                        processed_fallback_conditions = processed_fallback_conditions[len(claim_condition_prefix):].strip()
                    
                    if processed_fallback_conditions: # Ensure not empty after stripping
                        row_dict[column_names[2]] = processed_fallback_conditions
            
            table_data.append(row_dict)

        if table_data:
            eval_df = pd.DataFrame(table_data, columns=column_names)
            add_df_to_table(document, eval_df)
        else:
            add_styled_paragraph(document, "ไม่พบข้อมูลผลการประเมินที่สามารถแสดงในตารางได้", italic=True)
    else:
        add_styled_paragraph(document, "ไม่พบผลการประเมินคำกล่าวอ้าง", italic=True)
    document.add_paragraph() # Extra space after the entire evaluation section
    
    # 4. Disclaimers
    add_styled_heading(document, "ข้อความที่ต้องแสดงเพิ่มเติม (Disclaimers)", level=2, section_number="4.")
    disclaimer_results = report_data.get("disclaimer_results", [])
    if disclaimer_results:
        for disclaimer in disclaimer_results:
            p_msg = add_styled_paragraph(document, disclaimer.get('message', 'N/A'), color=COLOR_WARNING)
            details = f"   สารอาหาร: {disclaimer.get('nutrient')}, "
            if report_data.get("selected_label") != "ไม่อยู่ในบัญชีหมายเลข 2":
                 details += f"ค่าบนฉลาก: {disclaimer.get('label_value', 0):.1f} {disclaimer.get('unit')}, "
            details += f"ค่าจากหน่วยบริโภคอ้างอิง: {disclaimer.get('reference_value', 0):.1f} {disclaimer.get('unit')}, "
            details += f"ค่าที่กำหนด: {disclaimer.get('threshold', 0):.1f} {disclaimer.get('unit')}"
            p_details = add_styled_paragraph(document, details, color=COLOR_BLACK)
            p_details.paragraph_format.left_indent = Inches(0.25)
            document.add_paragraph()
    else:
        add_styled_paragraph(document, "ไม่พบข้อความ Disclaimer ที่ต้องแสดง", italic=True)

    add_page_numbers(document)

    # --- Global Font Override --- 
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = TARGET_FONT_NAME
            run.font.size = TARGET_FONT_SIZE
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = TARGET_FONT_NAME
                        run.font.size = TARGET_FONT_SIZE
    # --- End Global Font Override ---

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Example Usage (for testing - remove or comment out in final version)
if __name__ == '__main__':
    mock_report_data = {
        "selected_label": "ไม่อยู่ในบัญชีหมายเลข 2",
        "food_state_value": "solid",
        "nutrition_check_method": "ตรวจสอบจากฉลากโภชนาการ (ต่อ 1 หน่วยบริโภค)",
        "actual_serving_size": 50.0,
        "table_type": "table2",
        "has_added_sugar": None,
        "nutrient_inputs": {
            "energy": 200.0, "protein": 10.0, "fat": 8.0, 
            "saturated_fat": 1.0, "sugar": 5.0, "sodium": 100.0,
            "fiber": 3.0, "vitamin_a": 100.0, "vitamin_a_is_direct_rdi": False,
            "vitamin_c": 20.0
        },
        "RDI_MAPPING_ витамин": {
            "vitamin_a": "วิตามินเอ", "protein": "โปรตีน", "fat": "ไขมัน", 
            "saturated_fat": "ไขมันอิ่มตัว", "sugar": "น้ำตาล", "sodium": "โซเดียม", 
            "fiber": "ใยอาหาร", "energy": "พลังงาน", "vitamin_c": "วิตามินซี"
        },
        "VITAMIN_MINERAL_UNITS": {
            "vitamin_a": "µg RAE", "protein": "g", "fat": "g", "saturated_fat": "g", 
            "sugar": "g", "sodium": "mg", "fiber": "g", "energy": "kcal", "vitamin_c": "mg"
        },
        "REPORT_IGNORE_SUFFIXES": ["_rdi_percent", "_per_100kcal", "_energy_percent", "_is_direct_rdi", "_source"],
        "is_in_list_2": False,
        "adjusted_nutrient_values": { 
            "energy": 200.0, "protein": 10.0, "fat": 8.0, "saturated_fat": 1.0, 
            "sugar": 5.0, "sodium": 100.0, "fiber": 3.0, "vitamin_a": 100.0,
            "vitamin_c": 20.0, "protein_per_100kcal": 5.0, "fiber_per_100kcal": 1.5,
            "vitamin_a_rdi_percent": 12.5, "vitamin_c_rdi_percent": 33.3,
            "saturated_fat_energy_percent": 4.5
        },
        "rounded_values_display_df": pd.DataFrame({
            "สารอาหาร": ["พลังงาน", "โปรตีน", "วิตามินซี"], 
            "ค่าก่อนปัดเลข": [200, 10.0, 20.0], 
            "ค่าที่แสดงบนฉลาก": ["200", "10", "20"], 
            "หน่วย": ["kcal", "g", "mg"]
        }),
        "rdi_display_df": pd.DataFrame({
            "สารอาหาร": ["โปรตีน", "วิตามินเอ", "วิตามินซี"], 
            "%Thai RDI จากหน่วยบริโภคบนฉลาก": ["20%", "12.5%", "33.3%"], 
            "%Thai RDI ตามหน่วยบริโภคอ้างอิง": ["20%", "12.5%", "33.3%"]
        }),
        "saturated_fat_energy_percent": 4.5,
        "label_saturated_fat_energy_percent": 4.5, 
        "evaluation_messages": [
            {"text": "✅ โปรตีน: สามารถใช้คำกล่าวอ้าง: 'แหล่งของโปรตีน' (protein>=5g)", "is_success": True, "conditions_text": "This is a test condition for protein."}, 
            {"text": "❌ ไขมัน: ไม่เข้าเงื่อนไข 'ไขมันต่ำ' (fat<=3g)", "is_success": False},
            {"text": "⚠️ น้ำตาล: ควรระวังปริมาณน้ำตาล (sugar < 10g suggested)\n   📌 เงื่อนไขเพิ่มเติม: ตรวจสอบปริมาณน้ำตาลที่เติม", "is_success": False, "conditions_text": "Warning"}, 
            {"text": "✅ ใยอาหาร: สามารถใช้คำกล่าวอ้าง: 'แหล่งของใยอาหาร' (fiber>=1.5g) [ผ่านเงื่อนไขต่อ 100g/ml]", "is_success": True} 
        ],
        "disclaimer_results": [
            {
                "nutrient": "โซเดียม", "label_value": 50.0, "reference_value": 100.0, 
                "threshold": 120.0, "unit": "mg", 
                "message": "⚠️ ปริมาณ โซเดียม อยู่ในเกณฑ์ที่ต้องมีคำชี้แจง..."
            }
        ]
    }
    
    doc_stream = generate_nutrition_report(mock_report_data)
    with open("nutrition_report_formal_test.docx", "wb") as f:
        f.write(doc_stream.getvalue())
    print("Test report 'nutrition_report_formal_test.docx' generated.") 